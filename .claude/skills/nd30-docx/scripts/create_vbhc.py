#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo Văn bản Hành chính theo Nghị định 30/2020/NĐ-CP
Script sử dụng python-docx để tạo file .docx đúng thể thức ND30.

Chức năng:
1. create_vbhc()        — Tạo văn bản hành chính hoàn chỉnh
2. check_docx()         — Kiểm tra thể thức file .docx có đúng ND30 không
3. create_all_templates() — Tạo tất cả mẫu văn bản vào thư mục assets/
"""

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import sys


# ============================================================================
# HẰN SỐ ND30
# ============================================================================

# Lề trang (mm)
MARGIN_TOP = 20      # 20–25 mm
MARGIN_BOTTOM = 20   # 20–25 mm
MARGIN_LEFT = 30     # 30–35 mm
MARGIN_RIGHT = 15    # 15–20 mm

# Phông chữ
FONT_NAME = "Times New Roman"
FONT_COLOR = RGBColor(0, 0, 0)  # Đen

# Cỡ chữ mặc định (Phương án B: QH 13, TN 14, ND 14)
SIZE_QUOCHIEU = 13       # Quốc hiệu: 12–13
SIZE_TIEUUEU = 14        # Tiêu ngữ: 13–14
SIZE_COQUAN_CHUQUAN = 13 # Cơ quan chủ quản: 12–13
SIZE_COQUAN_BANHANH = 13 # Cơ quan ban hành: 12–13
SIZE_SO_KYHIEU = 13      # Số, ký hiệu: 13
SIZE_DIADANH = 14        # Địa danh + ngày: 13–14
SIZE_TEN_LOAI = 14       # Tên loại VB: 13–14
SIZE_TRICH_YEU = 14      # Trích yếu (có TL): 13–14
SIZE_TRICH_YEU_CV = 13   # Trích yếu (CV V/v): 12–13
SIZE_NOIDUNG = 14        # Nội dung: 13–14
SIZE_CANCU = 14          # Căn cứ: 13–14
SIZE_QUYEN_HAN = 14      # Quyền hạn ký: 13–14
SIZE_HOTEN = 14          # Họ tên ký: 13–14
SIZE_KINHUUI = 14        # Kính gửi: 13–14
SIZE_NOINHAN_LABEL = 12  # "Nơi nhận:": 12
SIZE_NOINHAN_LIST = 11   # DS nơi nhận: 11
SIZE_SOTRANG = 14        # Số trang: 13–14


# ============================================================================
# HÀM TIỆN ÍCH
# ============================================================================

def set_font(run, name=FONT_NAME, size=14, bold=False, italic=False, color=FONT_COLOR):
    """Thiết lập font cho một run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # Đảm bảo Times New Roman cho ký tự CJK/Unicode
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}" w:eastAsia="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:cs'), name)
        rFonts.set(qn('w:eastAsia'), name)


def set_paragraph_format(paragraph, alignment=None, space_before=0, space_after=0,
                         line_spacing=1.0, first_line_indent=None):
    """Thiết lập định dạng đoạn."""
    pf = paragraph.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)


def add_horizontal_line(cell_or_doc, width_ratio=1.0):
    """Thêm đường kẻ ngang bằng ký tự gạch dưới."""
    line_length = int(45 * width_ratio)
    p = cell_or_doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    run = p.add_run("_" * line_length)
    set_font(run, size=10, bold=False)
    return p


def remove_table_borders(table):
    """Xóa viền bảng (bảng ẩn)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def remove_cell_margins(table):
    """Xóa padding trong ô bảng."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    margins = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        '<w:top w:w="0" w:type="dxa"/>'
        '<w:left w:w="0" w:type="dxa"/>'
        '<w:bottom w:w="0" w:type="dxa"/>'
        '<w:right w:w="0" w:type="dxa"/>'
        '</w:tblCellMar>'
    )
    existing = tblPr.find(qn('w:tblCellMar'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(margins)


# ============================================================================
# KHỐI HEADER (Ô 1, 2, 3, 4, 5b)
# ============================================================================

def add_header_block(doc, co_quan_chu_quan, co_quan_ban_hanh,
                     so, ky_hieu, dia_danh, ngay, thang, nam,
                     trich_yeu_cv=None):
    """
    Tạo phần header bằng bảng 2 cột ẩn viền:
    - Cột trái: Cơ quan (Ô 2) + Số ký hiệu (Ô 3) + V/v (Ô 5b nếu CV)
    - Cột phải: Quốc hiệu + Tiêu ngữ (Ô 1) + Địa danh (Ô 4)
    """
    # Bảng 2 cột: cột trái ~45%, cột phải ~55%
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    remove_cell_margins(table)

    # Độ rộng cột (tính bằng EMU, tổng = page_width - margins)
    page_content_width = Mm(210) - Mm(MARGIN_LEFT) - Mm(MARGIN_RIGHT)  # 160mm
    col_left_width = int(page_content_width * 0.45)
    col_right_width = int(page_content_width * 0.55)
    for cell in table.columns[0].cells:
        cell.width = col_left_width
    for cell in table.columns[1].cells:
        cell.width = col_right_width

    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)

    # Xóa paragraph mặc định
    cell_left.paragraphs[0].clear()
    cell_right.paragraphs[0].clear()

    # ── CỘT TRÁI: Ô 2 (Cơ quan) ──

    # Dòng 1: Cơ quan chủ quản (IN HOA, KHÔNG đậm)
    p_chuquan = cell_left.paragraphs[0]
    set_paragraph_format(p_chuquan, WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.0)
    run = p_chuquan.add_run(co_quan_chu_quan.upper())
    set_font(run, size=SIZE_COQUAN_CHUQUAN, bold=False)

    # Dòng 2: Cơ quan ban hành (IN HOA, ĐẬM)
    p_banhanh = cell_left.add_paragraph()
    set_paragraph_format(p_banhanh, WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.0)
    run = p_banhanh.add_run(co_quan_ban_hanh.upper())
    set_font(run, size=SIZE_COQUAN_BANHANH, bold=True)

    # Đường kẻ dưới tên cơ quan ban hành (1/3 dòng chữ)
    p_line = cell_left.add_paragraph()
    set_paragraph_format(p_line, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=3)
    line_len = max(12, len(co_quan_ban_hanh) // 3)
    run = p_line.add_run("_" * line_len)
    set_font(run, size=10)

    # Ô 3: Số ký hiệu
    p_so = cell_left.add_paragraph()
    set_paragraph_format(p_so, WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=0)
    run = p_so.add_run("Số: ")
    set_font(run, size=SIZE_SO_KYHIEU, bold=False)
    run = p_so.add_run(f"{so}/{ky_hieu}")
    set_font(run, size=SIZE_SO_KYHIEU, bold=False)

    # Ô 5b: Trích yếu V/v (chỉ cho Công văn)
    if trich_yeu_cv:
        p_vv = cell_left.add_paragraph()
        set_paragraph_format(p_vv, WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=0)
        run = p_vv.add_run(f"V/v {trich_yeu_cv}")
        set_font(run, size=SIZE_TRICH_YEU_CV, bold=False)

    # ── CỘT PHẢI: Ô 1 (Quốc hiệu + Tiêu ngữ) ──

    # Dòng 1: Quốc hiệu (IN HOA, ĐẬM)
    p_qh = cell_right.paragraphs[0]
    set_paragraph_format(p_qh, WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.0)
    run = p_qh.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    set_font(run, size=SIZE_QUOCHIEU, bold=True)

    # Dòng 2: Tiêu ngữ (in thường, ĐẬM, ĐỨNG — KHÔNG nghiêng!)
    p_tn = cell_right.add_paragraph()
    set_paragraph_format(p_tn, WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.0)
    run = p_tn.add_run("Độc lập - Tự do - Hạnh phúc")
    set_font(run, size=SIZE_TIEUUEU, bold=True, italic=False)  # ⚠️ KHÔNG nghiêng

    # Đường kẻ dưới Tiêu ngữ (dài bằng dòng chữ)
    p_line = cell_right.add_paragraph()
    set_paragraph_format(p_line, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=3)
    run = p_line.add_run("_" * 30)
    set_font(run, size=10)

    # Ô 4: Địa danh + Ngày tháng năm (NGHIÊNG)
    p_dd = cell_right.add_paragraph()
    set_paragraph_format(p_dd, WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=0)
    # Đảm bảo ngày/tháng < 10 có số 0
    ngay_str = str(ngay).zfill(2)
    thang_str = str(thang).zfill(2)
    run = p_dd.add_run(f"{dia_danh}, ngày {ngay_str} tháng {thang_str} năm {nam}")
    set_font(run, size=SIZE_DIADANH, bold=False, italic=True)  # ✅ Nghiêng

    return table


# ============================================================================
# KHỐI TÊN LOẠI VÀ TRÍCH YẾU (Ô 5a)
# ============================================================================

def add_ten_loai_block(doc, ten_loai_vb, trich_yeu):
    """
    Thêm tên loại VB + trích yếu (Ô 5a) cho VB có tên loại.
    Canh giữa, phía dưới header block.
    """
    # Tên loại VB: IN HOA, ĐẬM, ĐỨNG
    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=0)
    run = p.add_run(ten_loai_vb.upper())
    set_font(run, size=SIZE_TEN_LOAI, bold=True)

    # Trích yếu: in thường, ĐẬM, ĐỨNG
    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    run = p.add_run(trich_yeu)
    set_font(run, size=SIZE_TRICH_YEU, bold=True)

    # Đường kẻ (1/3–1/2 dòng chữ)
    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)
    line_len = max(12, len(trich_yeu) // 3)
    run = p.add_run("_" * line_len)
    set_font(run, size=10)


# ============================================================================
# KHỐI KÍNH GỬI (Ô 9a)
# ============================================================================

def add_kinh_gui_block(doc, kinh_gui_list):
    """
    Thêm phần "Kính gửi:" cho Công văn, Tờ trình, Báo cáo.
    kinh_gui_list: list các nơi nhận kính gửi.
    """
    if not kinh_gui_list:
        return

    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6,
                         first_line_indent=1.0)

    if len(kinh_gui_list) == 1:
        # Gửi 1 nơi: trên cùng dòng
        run = p.add_run("Kính gửi: ")
        set_font(run, size=SIZE_KINHUUI, bold=False)
        run = p.add_run(kinh_gui_list[0])
        set_font(run, size=SIZE_KINHUUI, bold=False)
    else:
        # Gửi nhiều nơi: riêng dòng
        run = p.add_run("Kính gửi:")
        set_font(run, size=SIZE_KINHUUI, bold=False)
        for i, noi in enumerate(kinh_gui_list):
            p_noi = doc.add_paragraph()
            set_paragraph_format(p_noi, WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0)
            # Tabs/indent để thẳng hàng
            end_char = "." if i == len(kinh_gui_list) - 1 else ";"
            spaces = "                    "  # Khoảng cách thụt vào dưới dấu ":"
            run = p_noi.add_run(f"{spaces}- {noi}{end_char}")
            set_font(run, size=SIZE_KINHUUI, bold=False)


# ============================================================================
# KHỐI CĂN CỨ BAN HÀNH
# ============================================================================

def add_can_cu_block(doc, can_cu_list):
    """
    Thêm phần căn cứ ban hành (in thường, NGHIÊNG).
    can_cu_list: list các căn cứ.
    """
    if not can_cu_list:
        return

    for i, cc in enumerate(can_cu_list):
        p = doc.add_paragraph()
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=0,
                             first_line_indent=1.0)
        end_char = "." if i == len(can_cu_list) - 1 else ";"
        run = p.add_run(f"{cc}{end_char}")
        set_font(run, size=SIZE_CANCU, bold=False, italic=True)  # Nghiêng


# ============================================================================
# KHỐI NỘI DUNG (Ô 6)
# ============================================================================

def add_content_paragraph(doc, text, bold=False, italic=False, indent=True,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    """Thêm một đoạn nội dung văn bản."""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment, space_before=0, space_after=space_after,
                         line_spacing=1.15, first_line_indent=1.0 if indent else None)
    run = p.add_run(text)
    set_font(run, size=SIZE_NOIDUNG, bold=bold, italic=italic)
    return p


def add_dieu(doc, so_dieu, tieu_de):
    """Thêm Điều (đứng, đậm)."""
    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=6, space_after=6,
                         first_line_indent=1.0)
    run = p.add_run(f"Điều {so_dieu}. {tieu_de}")
    set_font(run, size=SIZE_NOIDUNG, bold=True)
    return p


# ============================================================================
# KHỐI KÝ TÊN (Ô 7, 8, 9b, 12)
# ============================================================================

def add_signature_block(doc, quyen_han_ky, chuc_vu, ho_ten, noi_nhan=None):
    """
    Tạo phần ký tên và nơi nhận bằng bảng 2 cột ẩn viền:
    - Cột trái: Nơi nhận (Ô 9b)
    - Cột phải: Quyền hạn, chức vụ, chữ ký, họ tên (Ô 7)
    """
    if noi_nhan is None:
        noi_nhan = ["Lưu: VT."]

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    remove_cell_margins(table)

    page_content_width = Mm(210) - Mm(MARGIN_LEFT) - Mm(MARGIN_RIGHT)
    col_left_width = int(page_content_width * 0.50)
    col_right_width = int(page_content_width * 0.50)
    for cell in table.columns[0].cells:
        cell.width = col_left_width
    for cell in table.columns[1].cells:
        cell.width = col_right_width

    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)

    cell_left.paragraphs[0].clear()
    cell_right.paragraphs[0].clear()

    # ── CỘT TRÁI: Ô 9b (Nơi nhận) ──

    # "Nơi nhận:" — nghiêng, đậm, cỡ 12
    p_nn = cell_left.paragraphs[0]
    set_paragraph_format(p_nn, WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0)
    run = p_nn.add_run("Nơi nhận:")
    set_font(run, size=SIZE_NOINHAN_LABEL, bold=True, italic=True)

    # Danh sách nơi nhận — đứng, cỡ 11
    for i, noi in enumerate(noi_nhan):
        p = cell_left.add_paragraph()
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0)
        is_last = (i == len(noi_nhan) - 1)
        if noi.startswith("Lưu:") or noi.startswith("- Lưu:"):
            text = f"- {noi}" if not noi.startswith("-") else noi
            if not text.endswith("."):
                text += "."
        else:
            text = f"- {noi}" if not noi.startswith("-") else noi
            if is_last and not text.endswith("."):
                text += "."
            elif not is_last and not text.endswith(";"):
                text += ";"
        run = p.add_run(text)
        set_font(run, size=SIZE_NOINHAN_LIST, bold=False)

    # ── CỘT PHẢI: Ô 7 (Ký tên) ──

    # Quyền hạn ký: IN HOA, ĐẬM
    p_qh = cell_right.paragraphs[0]
    set_paragraph_format(p_qh, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    run = p_qh.add_run(quyen_han_ky.upper())
    set_font(run, size=SIZE_QUYEN_HAN, bold=True)

    # Chức vụ: IN HOA, ĐẬM
    p_cv = cell_right.add_paragraph()
    set_paragraph_format(p_cv, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    run = p_cv.add_run(chuc_vu.upper())
    set_font(run, size=SIZE_QUYEN_HAN, bold=True)

    # Khoảng trống cho chữ ký (3 dòng trống)
    for _ in range(3):
        p = cell_right.add_paragraph()
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
        run = p.add_run("")
        set_font(run, size=SIZE_HOTEN)

    # Họ tên: in thường, ĐẬM, KHÔNG học hàm học vị
    p_ht = cell_right.add_paragraph()
    set_paragraph_format(p_ht, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    run = p_ht.add_run(ho_ten)
    set_font(run, size=SIZE_HOTEN, bold=True)

    return table


# ============================================================================
# HÀM CHÍNH: TẠO VĂN BẢN HÀNH CHÍNH
# ============================================================================

def create_vbhc(
    loai_vb,                        # 'QD', 'CV', 'TB', 'TTr', 'BC', 'KH', 'CT', 'HD'...
    co_quan_chu_quan,               # Tên cơ quan chủ quản
    co_quan_ban_hanh,               # Tên cơ quan ban hành
    so,                             # Số văn bản (str, vd: '05')
    ky_hieu,                        # Ký hiệu (vd: 'QĐ-UBND')
    dia_danh,                       # Địa danh (vd: 'Long Xuyên')
    ngay, thang, nam,               # Ngày, tháng, năm
    ten_loai_vb='',                 # Tên loại VB (vd: 'QUYẾT ĐỊNH')
    trich_yeu='',                   # Trích yếu nội dung
    noi_dung=None,                  # Nội dung văn bản (list hoặc str)
    can_cu=None,                    # Danh sách căn cứ (list)
    kinh_gui=None,                  # Danh sách kính gửi (list)
    quyen_han_ky='',                # Quyền hạn ký (vd: 'TM. ỦY BAN NHÂN DÂN')
    chuc_vu='',                     # Chức vụ ký (vd: 'GIÁM ĐỐC')
    ho_ten='',                      # Họ tên người ký
    noi_nhan=None,                  # Danh sách nơi nhận
    output_path='output.docx'       # Đường dẫn file đầu ra
):
    """Tạo file .docx văn bản hành chính hoàn chỉnh theo ND30."""

    doc = Document()

    # === THIẾT LẬP TRANG ===
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(MARGIN_TOP)
    section.bottom_margin = Mm(MARGIN_BOTTOM)
    section.left_margin = Mm(MARGIN_LEFT)
    section.right_margin = Mm(MARGIN_RIGHT)

    # Ẩn số trang đầu tiên
    section.different_first_page_header_footer = True

    # === HEADER BLOCK (Ô 1, 2, 3, 4, 5b) ===
    is_cong_van = (loai_vb.upper() == 'CV')
    trich_yeu_cv = trich_yeu if is_cong_van else None

    add_header_block(doc, co_quan_chu_quan, co_quan_ban_hanh,
                     so, ky_hieu, dia_danh, ngay, thang, nam,
                     trich_yeu_cv=trich_yeu_cv)

    # === TÊN LOẠI VÀ TRÍCH YẾU (Ô 5a) — chỉ cho VB có tên loại ===
    if not is_cong_van and ten_loai_vb:
        add_ten_loai_block(doc, ten_loai_vb, trich_yeu)

    # === CĂN CỨ BAN HÀNH (nếu có) ===
    if can_cu:
        add_can_cu_block(doc, can_cu)

    # === PHẦN QUYẾT ĐỊNH: / NỘI DUNG ===
    if loai_vb.upper() == 'QD' and ten_loai_vb:
        p = doc.add_paragraph()
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
        run = p.add_run("QUYẾT ĐỊNH:")
        set_font(run, size=SIZE_NOIDUNG, bold=True)

    # === KÍNH GỬI (Ô 9a) — cho CV, TTr, BC ===
    has_kinh_gui = loai_vb.upper() in ['CV', 'TTR', 'BC']
    if has_kinh_gui and kinh_gui:
        add_kinh_gui_block(doc, kinh_gui)

    # === NỘI DUNG VĂN BẢN (Ô 6) ===
    if noi_dung:
        if isinstance(noi_dung, str):
            noi_dung = [noi_dung]
        for nd in noi_dung:
            add_content_paragraph(doc, nd)

    # === KẾT THÚC VĂN BẢN: ./. ===
    p_end = doc.add_paragraph()
    set_paragraph_format(p_end, WD_ALIGN_PARAGRAPH.RIGHT, space_before=6, space_after=12)
    run = p_end.add_run("./.")
    set_font(run, size=SIZE_NOIDUNG, bold=False)

    # === KHỐI KÝ TÊN + NƠI NHẬN (Ô 7, 9b) ===
    if noi_nhan is None:
        noi_nhan = ["Lưu: VT."]

    add_signature_block(doc, quyen_han_ky, chuc_vu, ho_ten, noi_nhan)

    # === LƯU FILE ===
    doc.save(output_path)
    print(f"✅ Đã tạo văn bản: {output_path}")
    return output_path


# ============================================================================
# KIỂM TRA THỂ THỨC FILE .DOCX
# ============================================================================

def check_docx(file_path):
    """
    Kiểm tra cơ bản thể thức ND30 của file .docx.
    Trả về dict chứa kết quả kiểm tra.
    """
    results = {
        'file': file_path,
        'errors': [],
        'warnings': [],
        'passed': []
    }

    try:
        doc = Document(file_path)
    except Exception as e:
        results['errors'].append(f"Không thể mở file: {e}")
        return results

    section = doc.sections[0]

    # Kiểm tra lề
    left_mm = section.left_margin / Mm(1) if section.left_margin else 0
    right_mm = section.right_margin / Mm(1) if section.right_margin else 0
    top_mm = section.top_margin / Mm(1) if section.top_margin else 0
    bottom_mm = section.bottom_margin / Mm(1) if section.bottom_margin else 0

    # Dùng tolerance 0.5mm cho sai số floating-point
    TOL = 0.5
    if 30 - TOL <= left_mm <= 35 + TOL:
        results['passed'].append(f"Lề trái: {left_mm:.0f} mm ✅")
    else:
        results['errors'].append(f"Lề trái: {left_mm:.0f} mm ❌ (cần 30–35 mm)")

    if 15 - TOL <= right_mm <= 20 + TOL:
        results['passed'].append(f"Lề phải: {right_mm:.0f} mm ✅")
    else:
        results['errors'].append(f"Lề phải: {right_mm:.0f} mm ❌ (cần 15–20 mm)")

    if 20 - TOL <= top_mm <= 25 + TOL:
        results['passed'].append(f"Lề trên: {top_mm:.0f} mm ✅")
    else:
        results['errors'].append(f"Lề trên: {top_mm:.0f} mm ❌ (cần 20–25 mm)")

    if 20 - TOL <= bottom_mm <= 25 + TOL:
        results['passed'].append(f"Lề dưới: {bottom_mm:.0f} mm ✅")
    else:
        results['errors'].append(f"Lề dưới: {bottom_mm:.0f} mm ❌ (cần 20–25 mm)")

    # Kiểm tra khổ giấy
    w_mm = section.page_width / Mm(1) if section.page_width else 0
    h_mm = section.page_height / Mm(1) if section.page_height else 0
    if 209 <= w_mm <= 211 and 296 <= h_mm <= 298:
        results['passed'].append(f"Khổ giấy: {w_mm:.0f}×{h_mm:.0f} mm (A4) ✅")
    else:
        results['errors'].append(f"Khổ giấy: {w_mm:.0f}×{h_mm:.0f} mm ❌ (cần A4: 210×297)")

    # Kiểm tra font
    fonts_found = set()
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name:
                fonts_found.add(run.font.name)

    if fonts_found:
        non_tnr = fonts_found - {FONT_NAME}
        if not non_tnr:
            results['passed'].append(f"Font: Times New Roman ✅")
        else:
            results['errors'].append(
                f"Font không đúng: {', '.join(non_tnr)} ❌ (cần Times New Roman)"
            )
    else:
        results['warnings'].append("Không tìm thấy thông tin font (có thể dùng style mặc định)")

    # Kiểm tra nội dung cơ bản
    full_text = "\n".join([p.text for p in doc.paragraphs])
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    if "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in full_text:
        results['passed'].append("Quốc hiệu: có ✅")
    else:
        results['errors'].append("Quốc hiệu: thiếu ❌")

    if "Độc lập - Tự do - Hạnh phúc" in full_text:
        results['passed'].append("Tiêu ngữ: có ✅")
    elif "Độc lập" in full_text and "Tự do" in full_text:
        results['warnings'].append("Tiêu ngữ: có nhưng format có thể sai ⚠️")
    else:
        results['errors'].append("Tiêu ngữ: thiếu ❌")

    if "Số:" in full_text or "Số :" in full_text:
        results['passed'].append("Số ký hiệu: có ✅")
    else:
        results['warnings'].append("Số ký hiệu: không tìm thấy ⚠️")

    if "ngày" in full_text and "tháng" in full_text and "năm" in full_text:
        results['passed'].append("Ngày tháng năm: có ✅")
    else:
        results['warnings'].append("Ngày tháng năm: không tìm thấy ⚠️")

    # Tổng kết
    total = len(results['passed']) + len(results['errors']) + len(results['warnings'])
    print(f"\n📋 Kết quả kiểm tra: {file_path}")
    print(f"{'='*50}")
    for item in results['passed']:
        print(f"  ✅ {item}")
    for item in results['warnings']:
        print(f"  ⚠️ {item}")
    for item in results['errors']:
        print(f"  ❌ {item}")
    print(f"{'='*50}")
    print(f"Đạt: {len(results['passed'])}/{total} | Lỗi: {len(results['errors'])} | "
          f"Cảnh báo: {len(results['warnings'])}")

    return results


# ============================================================================
# TẠO TẤT CẢ MẪU VĂN BẢN
# ============================================================================

def create_all_templates(output_dir="assets"):
    """Tạo tất cả mẫu văn bản đúng chuẩn ND30 vào thư mục assets."""
    os.makedirs(output_dir, exist_ok=True)

    # 1. Mẫu Quyết định
    create_vbhc(
        loai_vb='QD',
        co_quan_chu_quan='UBND TỈNH AN GIANG',
        co_quan_ban_hanh='SỞ GIÁO DỤC VÀ ĐÀO TẠO',
        so='05', ky_hieu='QĐ-SGD&ĐT',
        dia_danh='Long Xuyên', ngay='15', thang='3', nam='2026',
        ten_loai_vb='QUYẾT ĐỊNH',
        trich_yeu='Về việc ban hành Quy chế làm việc của Sở Giáo dục và Đào tạo',
        can_cu=[
            'Căn cứ Luật Tổ chức chính quyền địa phương ngày 19 tháng 6 năm 2015',
            'Căn cứ Luật sửa đổi, bổ sung một số điều của Luật Tổ chức Chính phủ và Luật Tổ chức chính quyền địa phương ngày 22 tháng 11 năm 2019',
            'Căn cứ Nghị định số 30/2020/NĐ-CP ngày 05 tháng 3 năm 2020 của Chính phủ về công tác văn thư',
            'Theo đề nghị của Chánh Văn phòng Sở Giáo dục và Đào tạo'
        ],
        noi_dung=[
            'Điều 1. Ban hành kèm theo Quyết định này Quy chế làm việc của Sở Giáo dục và Đào tạo tỉnh An Giang.',
            'Điều 2. Quyết định này có hiệu lực thi hành kể từ ngày ký ban hành. Các quy định trước đây trái với Quyết định này đều bãi bỏ.',
            'Điều 3. Chánh Văn phòng, Trưởng các phòng chuyên môn, nghiệp vụ thuộc Sở và Thủ trưởng các đơn vị có liên quan chịu trách nhiệm thi hành Quyết định này.'
        ],
        quyen_han_ky='TM. ỦY BAN NHÂN DÂN',
        chuc_vu='GIÁM ĐỐC',
        ho_ten='Nguyễn Văn A',
        noi_nhan=[
            'UBND tỉnh (để báo cáo)',
            'Sở Nội vụ',
            'Các phòng, ban thuộc Sở',
            'Lưu: VT, VP'
        ],
        output_path=os.path.join(output_dir, 'mau_quyet_dinh.docx')
    )

    # 2. Mẫu Công văn
    create_vbhc(
        loai_vb='CV',
        co_quan_chu_quan='BỘ NỘI VỤ',
        co_quan_ban_hanh='CỤC VĂN THƯ VÀ LƯU TRỮ NHÀ NƯỚC',
        so='123', ky_hieu='VTLTNN-VP',
        dia_danh='Hà Nội', ngay='05', thang='01', nam='2026',
        trich_yeu='nâng bậc lương đối với công chức, viên chức năm 2026',
        kinh_gui=['Bộ Tài chính'],
        noi_dung=[
            'Thực hiện Nghị định số 204/2004/NĐ-CP ngày 14 tháng 12 năm 2004 của Chính phủ về chế độ tiền lương đối với cán bộ, công chức, viên chức và lực lượng vũ trang, Cục Văn thư và Lưu trữ nhà nước kính đề nghị Bộ Tài chính xem xét, giải quyết việc nâng bậc lương đối với công chức, viên chức đủ điều kiện trong năm 2026.',
            'Danh sách kèm theo gồm 15 (mười lăm) trường hợp đề nghị nâng bậc lương thường xuyên và 03 (ba) trường hợp đề nghị nâng bậc lương trước thời hạn.',
            'Cục Văn thư và Lưu trữ nhà nước kính đề nghị Bộ Tài chính xem xét, quyết định.'
        ],
        quyen_han_ky='KT. CỤC TRƯỞNG',
        chuc_vu='PHÓ CỤC TRƯỞNG',
        ho_ten='Trần Văn B',
        noi_nhan=[
            'Như trên',
            'Bộ Nội vụ (để báo cáo)',
            'Lưu: VT, TCCB'
        ],
        output_path=os.path.join(output_dir, 'mau_cong_van.docx')
    )

    # 3. Mẫu Tờ trình
    create_vbhc(
        loai_vb='TTr',
        co_quan_chu_quan='UBND TỈNH AN GIANG',
        co_quan_ban_hanh='SỞ TÀI CHÍNH',
        so='08', ky_hieu='TTr-STC',
        dia_danh='Long Xuyên', ngay='20', thang='3', nam='2026',
        ten_loai_vb='TỜ TRÌNH',
        trich_yeu='Về việc xin chủ trương đầu tư nâng cấp hệ thống công nghệ thông tin',
        kinh_gui=['Ủy ban nhân dân tỉnh An Giang'],
        noi_dung=[
            'I. SỰ CẦN THIẾT VÀ CƠ SỞ PHÁP LÝ',
            '1. Sự cần thiết: Hệ thống công nghệ thông tin hiện tại của Sở Tài chính đã sử dụng hơn 10 năm, nhiều thiết bị đã xuống cấp nghiêm trọng, không đáp ứng được yêu cầu chuyển đổi số của tỉnh.',
            '2. Cơ sở pháp lý: Căn cứ Quyết định số 749/QĐ-TTg ngày 03/6/2020 của Thủ tướng Chính phủ phê duyệt "Chương trình Chuyển đổi số quốc gia đến năm 2025, định hướng đến năm 2030".',
            'II. NỘI DUNG ĐỀ XUẤT',
            '1. Nâng cấp hệ thống máy chủ và thiết bị lưu trữ dữ liệu.',
            '2. Triển khai hệ thống quản lý văn bản điện tử liên thông.',
            '3. Đào tạo, bồi dưỡng nguồn nhân lực công nghệ thông tin.',
            'III. KINH PHÍ DỰ KIẾN',
            'Tổng kinh phí dự kiến: 2.500.000.000 đồng (Hai tỷ năm trăm triệu đồng) từ nguồn ngân sách tỉnh.',
            'Kính trình Ủy ban nhân dân tỉnh xem xét, quyết định.'
        ],
        quyen_han_ky='',
        chuc_vu='GIÁM ĐỐC',
        ho_ten='Lê Văn C',
        noi_nhan=[
            'Như trên',
            'Sở Kế hoạch và Đầu tư (để phối hợp)',
            'Lưu: VT, KHTC'
        ],
        output_path=os.path.join(output_dir, 'mau_to_trinh.docx')
    )

    # 4. Mẫu Báo cáo
    create_vbhc(
        loai_vb='BC',
        co_quan_chu_quan='UBND TỈNH AN GIANG',
        co_quan_ban_hanh='SỞ GIÁO DỤC VÀ ĐÀO TẠO',
        so='15', ky_hieu='BC-SGD&ĐT',
        dia_danh='Long Xuyên', ngay='25', thang='3', nam='2026',
        ten_loai_vb='BÁO CÁO',
        trich_yeu='Tình hình thực hiện nhiệm vụ giáo dục quý I năm 2026',
        kinh_gui=['Ủy ban nhân dân tỉnh An Giang'],
        noi_dung=[
            'I. TÌNH HÌNH THỰC HIỆN NHIỆM VỤ',
            '1. Kết quả đạt được:',
            'a) Về giáo dục mầm non: Tỷ lệ huy động trẻ đạt 95%, vượt kế hoạch đề ra.',
            'b) Về giáo dục phổ thông: 100% học sinh được đến trường đúng độ tuổi.',
            '2. Tồn tại, hạn chế:',
            'Cơ sở vật chất một số trường học vùng sâu, vùng xa còn thiếu thốn.',
            'II. PHƯƠNG HƯỚNG, NHIỆM VỤ QUÝ II',
            '1. Tiếp tục triển khai chương trình giáo dục phổ thông mới.',
            '2. Đẩy mạnh ứng dụng công nghệ thông tin trong dạy và học.',
            'Trên đây là Báo cáo tình hình thực hiện nhiệm vụ giáo dục quý I năm 2026, kính trình Ủy ban nhân dân tỉnh xem xét.'
        ],
        quyen_han_ky='',
        chuc_vu='GIÁM ĐỐC',
        ho_ten='Nguyễn Văn A',
        noi_nhan=[
            'Như trên',
            'Bộ GD&ĐT (để báo cáo)',
            'Các phòng, ban thuộc Sở',
            'Lưu: VT, VP'
        ],
        output_path=os.path.join(output_dir, 'mau_bao_cao.docx')
    )

    # 5. Mẫu Thông báo
    create_vbhc(
        loai_vb='TB',
        co_quan_chu_quan='UBND TỈNH AN GIANG',
        co_quan_ban_hanh='SỞ GIÁO DỤC VÀ ĐÀO TẠO',
        so='22', ky_hieu='TB-SGD&ĐT',
        dia_danh='Long Xuyên', ngay='01', thang='4', nam='2026',
        ten_loai_vb='THÔNG BÁO',
        trich_yeu='Về việc tổ chức Hội nghị tổng kết năm học 2025-2026',
        noi_dung=[
            'Sở Giáo dục và Đào tạo thông báo kế hoạch tổ chức Hội nghị tổng kết năm học 2025-2026 như sau:',
            '1. Thời gian: 08 giờ 00, ngày 15 tháng 7 năm 2026 (Thứ Tư).',
            '2. Địa điểm: Hội trường lớn Sở Giáo dục và Đào tạo tỉnh An Giang, số 01 Lê Triệu Kiết, phường Mỹ Bình, thành phố Long Xuyên.',
            '3. Thành phần tham dự:',
            'a) Lãnh đạo Sở Giáo dục và Đào tạo.',
            'b) Trưởng, Phó các phòng, ban thuộc Sở.',
            'c) Trưởng phòng Giáo dục và Đào tạo các huyện, thị xã, thành phố.',
            'd) Hiệu trưởng các trường THPT, Trung tâm GDTX trên địa bàn tỉnh.',
            '4. Nội dung: Tổng kết đánh giá kết quả năm học 2025-2026; triển khai phương hướng, nhiệm vụ năm học 2026-2027.',
            'Đề nghị các đơn vị cử người tham dự đúng thành phần, thời gian.'
        ],
        quyen_han_ky='TL. GIÁM ĐỐC',
        chuc_vu='CHÁNH VĂN PHÒNG',
        ho_ten='Phạm Thị D',
        noi_nhan=[
            'Lãnh đạo Sở',
            'Các phòng, ban thuộc Sở',
            'Phòng GD&ĐT các huyện, TX, TP',
            'Các trường THPT, TT GDTX',
            'Lưu: VT, VP'
        ],
        output_path=os.path.join(output_dir, 'mau_thong_bao.docx')
    )

    # 6. Mẫu Kế hoạch
    create_vbhc(
        loai_vb='KH',
        co_quan_chu_quan='UBND TỈNH AN GIANG',
        co_quan_ban_hanh='SỞ GIÁO DỤC VÀ ĐÀO TẠO',
        so='30', ky_hieu='KH-SGD&ĐT',
        dia_danh='Long Xuyên', ngay='10', thang='4', nam='2026',
        ten_loai_vb='KẾ HOẠCH',
        trich_yeu='Tổ chức tập huấn chuyên môn cho giáo viên năm 2026',
        noi_dung=[
            'Thực hiện Kế hoạch công tác năm 2026, Sở Giáo dục và Đào tạo xây dựng Kế hoạch tổ chức tập huấn chuyên môn cho giáo viên như sau:',
            'I. MỤC ĐÍCH, YÊU CẦU',
            '1. Mục đích: Nâng cao năng lực chuyên môn, nghiệp vụ cho đội ngũ giáo viên trên địa bàn tỉnh, đáp ứng yêu cầu đổi mới giáo dục.',
            '2. Yêu cầu: Đảm bảo 100% giáo viên các cấp được tập huấn; nội dung thiết thực, hiệu quả.',
            'II. NỘI DUNG TẬP HUẤN',
            '1. Phương pháp dạy học tích cực theo chương trình giáo dục phổ thông mới.',
            '2. Ứng dụng công nghệ thông tin và chuyển đổi số trong dạy học.',
            '3. Kỹ năng đánh giá năng lực học sinh theo chuẩn đầu ra.',
            'III. THỜI GIAN, ĐỊA ĐIỂM',
            '1. Thời gian: Từ ngày 01/6/2026 đến ngày 30/6/2026.',
            '2. Địa điểm: Tại các trung tâm tập huấn cấp huyện.',
            'IV. KINH PHÍ',
            'Kinh phí thực hiện từ nguồn ngân sách sự nghiệp giáo dục năm 2026.',
            'V. TỔ CHỨC THỰC HIỆN',
            '1. Phòng Giáo dục Trung học: Chủ trì xây dựng nội dung tập huấn.',
            '2. Văn phòng Sở: Đảm bảo cơ sở vật chất, hậu cần.',
            '3. Phòng GD&ĐT các huyện: Tổ chức triển khai tại địa phương.'
        ],
        quyen_han_ky='',
        chuc_vu='GIÁM ĐỐC',
        ho_ten='Nguyễn Văn A',
        noi_nhan=[
            'UBND tỉnh (để báo cáo)',
            'Các phòng, ban thuộc Sở',
            'Phòng GD&ĐT các huyện, TX, TP',
            'Lưu: VT, GDTrH'
        ],
        output_path=os.path.join(output_dir, 'mau_ke_hoach.docx')
    )

    print(f"\n🎉 Đã tạo 6 mẫu văn bản trong thư mục '{output_dir}':")
    print("  1. mau_quyet_dinh.docx")
    print("  2. mau_cong_van.docx")
    print("  3. mau_to_trinh.docx")
    print("  4. mau_bao_cao.docx")
    print("  5. mau_thong_bao.docx")
    print("  6. mau_ke_hoach.docx")


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    if len(sys.argv) > 1:
        if sys.argv[1] == '--templates':
            # Tạo tất cả mẫu
            output = sys.argv[2] if len(sys.argv) > 2 else 'assets'
            create_all_templates(output)
        elif sys.argv[1] == '--check':
            # Kiểm tra file .docx
            if len(sys.argv) > 2:
                check_docx(sys.argv[2])
            else:
                print("Cách dùng: python create_vbhc.py --check <file.docx>")
        else:
            print("Cách dùng:")
            print("  python create_vbhc.py --templates [output_dir]  # Tạo mẫu")
            print("  python create_vbhc.py --check <file.docx>       # Kiểm tra")
    else:
        print("Cách dùng:")
        print("  python create_vbhc.py --templates [output_dir]  # Tạo mẫu")
        print("  python create_vbhc.py --check <file.docx>       # Kiểm tra")
