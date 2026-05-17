import sys
import json
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_style(run, font_name='Times New Roman', font_size=14, bold=False, italic=False):
    run.font.name = font_name
    # Ensure proper rendering of Times New Roman in docx internally
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic

def create_nd30_docx(json_path, output_path):
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    
    # 1. Page settings (A4, Margins: Top 2cm, Bottom 2cm, Left 3cm, Right 2cm)
    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    # Change default normal style to Times New Roman 14
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # 2. Header Table (Borderless, 2 columns)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(6.0)
    table.columns[1].width = Cm(10.0)

    cell_1 = table.cell(0, 0)
    cell_2 = table.cell(0, 1)

    # Column 1: Tên cơ quan chủ quản & Tên cơ quan ban hành
    p1 = cell_1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if data.get('co_quan_chu_quan'):
        run = p1.add_run(data['co_quan_chu_quan'].upper() + '\n')
        set_style(run, font_size=12)
    
    if data.get('co_quan_ban_hanh'):
        run = p1.add_run(data['co_quan_ban_hanh'].upper() + '\n')
        set_style(run, font_size=13, bold=True)
    
    # Dấu gạch ngang dưới tên cơ quan ban hành
    run = p1.add_run('-------\n')
    set_style(run, font_size=12)

    if data.get('so_ky_hieu'):
        run = p1.add_run(data['so_ky_hieu'])
        set_style(run, font_size=13)

    # Column 2: Quốc hiệu - Tiêu ngữ
    p2 = cell_2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n')
    set_style(run, font_size=12, bold=True)
    
    run = p2.add_run('Độc lập - Tự do - Hạnh phúc\n')
    set_style(run, font_size=13, bold=True)
    
    # Dấu gạch ngang dưới tiêu ngữ
    run = p2.add_run('---------------\n')
    set_style(run, font_size=12)

    if data.get('dia_danh_ngay_thang'):
        # Add some spacing
        p2.add_run('\n')
        run = p2.add_run(data['dia_danh_ngay_thang'])
        set_style(run, font_size=14, italic=True)

    # 3. Loại văn bản và trích yếu
    doc.add_paragraph() # Spacing
    
    if data.get('ten_loai_van_ban'):
        p_type = doc.add_paragraph()
        p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_type.add_run(data['ten_loai_van_ban'].upper())
        set_style(run, font_size=14, bold=True)
    
    if data.get('trich_yeu'):
        p_trich_yeu = doc.add_paragraph()
        p_trich_yeu.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_trich_yeu.add_run(data['trich_yeu'])
        set_style(run, font_size=14, bold=True)

    if data.get('kinh_gui'):
        p_kg = doc.add_paragraph()
        p_kg.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_kg.add_run(data['kinh_gui'])
        set_style(run, font_size=14)

    doc.add_paragraph() # Spacing

    # 4. Căn cứ
    if data.get('can_cu'):
        for cc in data['can_cu']:
            p_cc = doc.add_paragraph()
            p_cc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_cc.paragraph_format.first_line_indent = Cm(1.27)
            p_cc.paragraph_format.line_spacing = 1.5
            run = p_cc.add_run(cc)
            set_style(run, font_size=14, italic=True)

    # 5. Nội dung
    if data.get('noi_dung'):
        for block in data['noi_dung']:
            p_nd = doc.add_paragraph()
            p_nd.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_nd.paragraph_format.first_line_indent = Cm(1.27)
            p_nd.paragraph_format.line_spacing = 1.5
            p_nd.paragraph_format.space_before = Pt(6)
            
            if block.get('loai') == 'dieu':
                run = p_nd.add_run(block.get('tieu_de', '') + ' ')
                set_style(run, font_size=14, bold=True)
                run = p_nd.add_run(block.get('noi_dung', ''))
                set_style(run, font_size=14)
            else:
                run = p_nd.add_run(block.get('noi_dung', ''))
                set_style(run, font_size=14)

    doc.add_paragraph() # Spacing before signature

    # 6. Footer Table (Nơi nhận & Chữ ký)
    table_footer = doc.add_table(rows=1, cols=2)
    table_footer.autofit = False
    table_footer.columns[0].width = Cm(8.0)
    table_footer.columns[1].width = Cm(8.0)

    f_cell_1 = table_footer.cell(0, 0)
    f_cell_2 = table_footer.cell(0, 1)

    # Nơi nhận
    p_nn = f_cell_1.paragraphs[0]
    p_nn.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if data.get('noi_nhan'):
        run = p_nn.add_run('Nơi nhận:\n')
        set_style(run, font_size=12, bold=True, italic=True)
        for item in data['noi_nhan']:
            run = p_nn.add_run('- ' + item + '\n')
            set_style(run, font_size=11)

    # Chữ ký
    p_ck = f_cell_2.paragraphs[0]
    p_ck.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if data.get('chuc_danh_nguoi_ky'):
        run = p_ck.add_run(data['chuc_danh_nguoi_ky'].upper() + '\n')
        set_style(run, font_size=14, bold=True)
    
    # Khoảng trống cho chữ ký
    p_ck.add_run('\n\n\n\n')

    if data.get('ten_nguoi_ky'):
        run = p_ck.add_run(data['ten_nguoi_ky'])
        set_style(run, font_size=14, bold=True)

    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python generate_nd30_docx.py <input_json> <output_docx>")
        sys.exit(1)
    
    create_nd30_docx(sys.argv[1], sys.argv[2])
