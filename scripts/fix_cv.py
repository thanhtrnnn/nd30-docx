#!/usr/bin/env python3
"""Fix CV: Calibri to Times New Roman, left margin 25mm to 30mm."""
from docx import Document
from docx.shared import Mm
from docx.oxml.ns import qn

input_path = 'output/CV_PhamTuanAnh_AIEngineerIntern.docx'
output_path = 'output/CV_PhamTuanAnh_AIEngineerIntern_Rewritten.docx'

doc = Document(input_path)

# Fix margins
for section in doc.sections:
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)

# Fix font for all runs
def set_font(run):
    run.font.name = 'Times New Roman'
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:eastAsia'), 'Times New Roman')

for para in doc.paragraphs:
    for run in para.runs:
        set_font(run)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run)

doc.save(output_path)
print(f'Da tao: {output_path}')
