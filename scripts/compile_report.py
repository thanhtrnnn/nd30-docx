#!/usr/bin/env python3
"""Compile tất cả bài thi SKD1103 vào 1 file .docx duy nhất"""
import os
from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')

def set_font(paragraph, name='Times New Roman', size=Pt(13), bold=False, italic=False):
    """Set font cho paragraph"""
    for run in paragraph.runs:
        run.font.name = name
        run.font.size = size
        run.bold = bold
        run.italic = italic
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:eastAsia'), name)

def add_heading_styled(doc, text, level=1):
    """Thêm heading với font TNR"""
    p = doc.add_heading(text, level=level)
    set_font(p, size=Pt(14) if level == 1 else Pt(13), bold=True)
    return p

def add_para(doc, text, bold=False, italic=False, size=Pt(13), align=None):
    """Thêm paragraph với font TNR"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.bold = bold
    run.italic = italic
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p

def add_page_break(doc):
    doc.add_page_break()

def create_cover_page(doc):
    """Tạo trang bìa"""
    # Khoảng cách
    for _ in range(3):
        doc.add_paragraph()

    # Tên trường
    p = add_para(doc, 'HỌC VIỆN CÔNG NGHỆ', bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)
    p = add_para(doc, 'BƯU CHÍNH VIỄN THÔNG', bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # Tên khoa/viện
    p = add_para(doc, 'VIỆN KINH TẾ BƯU ĐIỆN', bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tên đề thi
    p = add_para(doc, 'BÁO CÁO CUỐI KỲ', bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    p = add_para(doc, 'HỌC PHẦN: KỸ NĂNG TẠO LẬP VĂN BẢN TIẾNG VIỆT', bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)
    p = add_para(doc, 'HỆ ĐẠI HỌC: CHÍNH QUY', bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tình huống
    p = add_para(doc, 'Đề tài:', bold=True, size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER)
    p = add_para(doc, 'Nền tảng AI hỗ trợ bảo tồn và giảng dạy tiếng Jrai', bold=True, size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER)
    p = add_para(doc, '(Mở rộng sang tiếng Bahnar)', size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER)
    p = add_para(doc, 'cho chuyển đổi số tỉnh Gia Lai', size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(3):
        doc.add_paragraph()

    # Thông tin nhóm
    p = add_para(doc, 'Nhóm: ...', size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER)
    p = add_para(doc, 'Lớp: ...', size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER)
    p = add_para(doc, 'Giảng viên hướng dẫn: ...', size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    p = add_para(doc, 'Hà Nội, tháng 5 năm 2026', size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER)

def create_toc(doc):
    """Tạo mục lục"""
    add_heading_styled(doc, 'MỤC LỤC', level=1)

    toc_items = [
        ('Lời mở đầu', '3'),
        ('PHẦN I: HỆ THỐNG LÝ THUYẾT', '4'),
        ('  Câu 1: Các thành phần thể thức chính của văn bản hành chính', '4'),
        ('  Câu 2: Yêu cầu về CV và lưu ý khi soạn thảo email ứng tuyển', '6'),
        ('PHẦN II: THỰC HÀNH SOẠN THẢO', '8'),
        ('  Nhiệm vụ 1: Bộ 03 văn bản hành chính', '8'),
        ('    1.1. Tờ trình', '8'),
        ('    1.2. Công văn', '9'),
        ('    1.3. Thông báo', '10'),
        ('  Nhiệm vụ 2: CV cá nhân', '11'),
        ('  Nhiệm vụ 3: Thư ứng tuyển và Email', '12'),
        ('Kết luận', '14'),
        ('Tài liệu tham khảo', '15'),
        ('Bảng đánh giá đóng góp', '16'),
        ('Slide thuyết trình', '17'),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        # Thêm tab và số trang
        run2 = p.add_run(f'\t{page}')
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(13)

def create_foreword(doc):
    """Tạo lời mở đầu"""
    add_heading_styled(doc, 'LỜI MỞ ĐẦU', level=1)

    add_para(doc,
        'Trong bối cảnh chuyển đổi số đang diễn ra mạnh mẽ trên toàn quốc, việc nắm vững '
        'các kỹ năng tạo lập văn bản tiếng Việt là yêu cầu thiết yếu đối với sinh viên đại học. '
        'Học phần Kỹ năng tạo lập văn bản tiếng Việt giúp sinh viên hiểu rõ các quy định '
        'về thể thức và kỹ thuật trình bày văn bản hành chính theo Nghị định 30/2020/NĐ-CP, '
        'đồng thời phát triển kỹ năng soạn thảo CV, thư ứng tuyển và email chuyên nghiệp.'
    )

    add_para(doc,
        'Báo cáo này được thực hiện dựa trên tình huống giả định về dự án "Nền tảng AI '
        'hỗ trợ bảo tồn và giảng dạy tiếng Jrai (mở rộng sang tiếng Bahnar) cho chuyển đổi số '
        'tỉnh Gia Lai" tại Học viện Công nghệ Bưu chính Viễn thông. Nội dung báo cáo bao gồm '
        'hai phần chính: hệ thống lý thuyết về các thành phần thể thức văn bản và yêu cầu CV/email, '
        'cùng với thực hành soạn thảo bộ văn bản hành chính, CV cá nhân, thư ứng tuyển và email.'
    )

    add_para(doc,
        'Báo cáo được thực hiện trong thời gian 10 ngày (từ ngày 07/5/2026 đến ngày 17/5/2026). '
        'Nhóm xin chân thành cảm ơn sự hướng dẫn của giảng viên và sự hỗ trợ của các thành viên '
        'trong nhóm trong quá trình thực hiện.'
    )

def create_theory(doc):
    """Tạo phần lý thuyết — phiên bản chi tiết với trích dẫn điều khoản NĐ30"""
    add_heading_styled(doc, 'PHẦN I: HỆ THỐNG LÝ THUYẾT', level=1)

    # ==================== CÂU 1 ====================
    add_heading_styled(doc, 'Câu 1: Chín thành phần thể thức văn bản hành chính theo NĐ30/2020/NĐ-CP', level=2)

    add_para(doc,
        'Theo Phụ lục I của Nghị định 30/2020/NĐ-CP ngày 05/3/2020 của Chính phủ, '
        'văn bản hành chính phải đảm bảo chín thành phần thể thức chính. Mỗi thành phần '
        'có quy định cụ thể về vị trí, cỡ chữ, kiểu chữ và cách trình bày.'
    )

    # 1. Quốc hiệu, tiêu ngữ
    add_para(doc, '1. Quốc hiệu và Tiêu ngữ', bold=True)
    add_para(doc,
        'Quốc hiệu "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" là thành phần bắt buộc, '
        'thể hiện chủ quyền quốc gia. Theo Khoản 1 Điều 5 NĐ30, quốc hiệu được trình bày '
        'bằng chữ IN HOA, cỡ chữ 12-13pt, kiểu đứng, đậm, đặt bên phải trên cùng trang đầu tiên. '
        'Tiêu ngữ "Độc lập - Tự do - Hạnh phúc" đặt ngay dưới quốc hiệu, cỡ chữ 13-14pt, '
        'in thường (chữ cái đầu mỗi cụm từ viết hoa), đứng, đậm, KHÔNG in nghiêng. '
        'Giữa hai dòng này là đường kẻ ngang nét liền, dài bằng dòng chữ. Theo Quy tắc thống nhất '
        'cỡ chữ (Mục I.5 Phụ lục I), nếu quốc hiệu dùng 12pt thì tiêu ngữ dùng 13pt: '
        'nếu quốc hiệu dùng 13pt thì tiêu ngữ dùng 14pt.'
    )

    # 2. Tên cơ quan chủ quản
    add_para(doc, '2. Tên cơ quan, tổ chức chủ quản', bold=True)
    add_para(doc,
        'Tên cơ quan chủ quản trực tiếp là dòng trên cùng bên trái của bảng header, '
        'trình bày bằng chữ IN HOA, cỡ 12-13pt, đứng, KHÔNG đậm. Nếu văn bản không có '
        'cơ quan chủ quản thì bỏ trống dòng này. Theo Khoản 2 Điều 5 NĐ30, thành phần này '
        'giúp xác định hệ thống hành chính mà cơ quan ban hành trực thuộc.'
    )

    # 3. Tên cơ quan ban hành
    add_para(doc, '3. Tên cơ quan, tổ chức ban hành văn bản', bold=True)
    add_para(doc,
        'Tên cơ quan ban hành đặt ngay dưới tên cơ quan chủ quản, trình bày bằng chữ IN HOA, '
        'cỡ 12-13pt, đứng, đậm. Phía dưới có đường kẻ ngang nét liền, dài từ 1/3 đến 1/2 dòng chữ, '
        'cân đối giữa dòng. Theo Khoản 3 Điều 5 NĐ30, đây là cơ quan có thẩm quyền ban hành văn bản, '
        'chịu trách nhiệm về nội dung. Hai dòng tên cơ quan cách nhau dòng đơn.'
    )

    # 4. Số, ký hiệu
    add_para(doc, '4. Số, ký hiệu của văn bản', bold=True)
    add_para(doc,
        'Số và ký hiệu văn bản là mã định danh, bao gồm số thứ tự và ký hiệu phân loại. '
        'Theo Khoản 4 Điều 5 NĐ30, từ "Số" in thường, cỡ 13pt, đứng, sau có dấu hai chấm. '
        'Số nhỏ hơn 10 phải thêm số 0 phía trước (01, 02... 09). Ký hiệu dùng chữ IN HOA, '
        'cỡ 13pt, ngăn cách bằng dấu gạch chéo và gạch nối. Ví dụ: "Số: 01/2025/TTr-PTIT". '
        'Thành phần này đặt bên trái bảng header, cùng hàng ngang với địa danh, ngày tháng.'
    )

    # 5. Địa danh, ngày tháng
    add_para(doc, '5. Địa danh và thời gian ban hành', bold=True)
    add_para(doc,
        'Địa danh và ngày tháng ban hành đặt bên phải bảng header, cùng hàng ngang với số ký hiệu. '
        'Theo Khoản 5 Điều 5 NĐ30, đây là thành phần DUY NHẤT được in nghiêng trong toàn bộ văn bản. '
        'Cỡ chữ 13-14pt, chữ cái đầu địa danh viết hoa, sau địa danh có dấu phẩy. '
        'Ngày và tháng nhỏ hơn 10 phải thêm số 0 phía trước (ví dụ: "ngày 05", "tháng 01").'
    )

    # 6. Tên loại và trích yếu
    add_para(doc, '6. Tên loại và trích yếu nội dung', bold=True)
    add_para(doc,
        'Tên loại văn bản (Quyết định, Tờ trình, Thông báo...) trình bày bằng chữ IN HOA, '
        'cỡ 13-14pt, đứng, đậm, căn giữa. Trích yếu nội dung đặt ngay dưới tên loại, in thường, '
        'cỡ 13-14pt, đứng, đậm, căn giữa. Phía dưới có đường kẻ ngang dài 1/3 đến 1/2 dòng chữ. '
        'Đối với công văn (không có tên loại), chỉ có trích yếu dạng "V/v...", cỡ 12-13pt, '
        'đặt dưới số ký hiệu, cách 6pt. Theo Khoản 6 Điều 5 NĐ30, trích yếu phải ngắn gọn, '
        'chính xác, phản ánh đúng nội dung chính của văn bản.'
    )

    # 7. Nội dung văn bản
    add_para(doc, '7. Nội dung văn bản', bold=True)
    add_para(doc,
        'Nội dung là phần chính của văn bản, trình bày bằng chữ in thường, đứng, cỡ 13-14pt, '
        'căn đều hai lề, lùi đầu dòng 1cm hoặc 1,27cm. Theo Khoản 7 Điều 5 NĐ30, khoảng cách '
        'giữa các đoạn tối thiểu 6pt, khoảng cách dòng tối thiểu dòng đơn, tối đa 1,5 lines. '
        'Căn cứ ban hành được in nghiêng, mỗi căn cứ trên một dòng riêng, cuối dòng dấu chấm phẩy, '
        'dòng cuối dấu chấm. Bố cục nội dung chia theo Phần, Chương, Mục, Điều, Khoản, Điểm '
        'với quy tắc đánh số và trình bày cụ thể.'
    )

    # 8. Chức vụ, họ tên, chữ ký
    add_para(doc, '8. Chức vụ, họ tên và chữ ký người có thẩm quyền', bold=True)
    add_para(doc,
        'Theo Khoản 8 Điều 5 NĐ30, quyền hạn ký (TM., KT., Q., TL., TUQ.) và chức vụ người ký '
        'trình bày bằng chữ IN HOA, cỡ 13-14pt, đứng, đậm, căn giữa bên phải. '
        'Dòng "(Ký, ghi rõ họ tên)" in nghiêng, cỡ 13-14pt, đặt giữa chức vụ và họ tên. '
        'Họ tên người ký in thường, đứng, đậm, cỡ 13-14pt, KHÔNG ghi học hàm, học vị. '
        'Chữ ký tay đặt giữa chức vụ và họ tên.'
    )

    # 9. Dấu cơ quan
    add_para(doc, '9. Dấu của cơ quan, tổ chức', bold=True)
    add_para(doc,
        'Dấu cơ quan là con dấu đỏ, được đóng trùm lên khoảng 1/3 chữ ký về bên trái, '
        'theo Khoản 9 Điều 5 NĐ30. Đối với văn bản điện tử, có thể sử dụng chữ ký số. '
        'Dấu chỉ có giá trị khi được đóng đúng vị trí và trên văn bản có đủ các thành phần thể thức. '
        'Không phải tất cả văn bản đều bắt buộc có dấu — tùy thuộc vào loại văn bản '
        'và quy định của cơ quan ban hành.'
    )

    # ==================== CÂU 2 ====================
    add_heading_styled(doc, 'Câu 2: Yêu cầu về CV và lưu ý khi viết email xin việc', level=2)

    add_heading_styled(doc, 'I. Curriculum Vitae (CV)', level=2)

    add_para(doc, '1. Định nghĩa và vai trò', bold=True)
    add_para(doc,
        'CV (Curriculum Vitae) là bản tóm tắt thông tin cá nhân, học vấn, kinh nghiệm, '
        'kỹ năng và thành tích của người ứng tuyển. CV đóng vai trò là "bộ mặt đầu tiên" '
        'trước nhà tuyển dụng, giúp họ đánh giá sự phù hợp của ứng viên với vị trí tuyển dụng. '
        'Theo giáo trình SKD1103, CV phải ngắn gọn (tối đa 2 trang A4), rõ ràng, chuyên nghiệp '
        'và được tối ưu cho hệ thống ATS (Applicant Tracking System).'
    )

    add_para(doc, '2. Mười một phần của CV', bold=True)

    cv_parts = [
        ('Tiêu đề CV', 'Gồm họ tên ứng viên và vị trí ứng tuyển. Nên nhúng keyword từ JD vào tiêu đề để ATS dễ nhận diện.'),
        ('Thông tin cá nhân', 'Bao gồm họ tên (in hoa, bold), email chuyên nghiệp (không nickname), số điện thoại đầy đủ, địa chỉ (chỉ ghi thành phố), và link LinkedIn/portfolio nếu có.'),
        ('Mục tiêu nghề nghiệp', 'Chia ngắn hạn (1-3 năm) và dài hạn (3-5 năm). Phải cụ thể, có keyword từ JD, nhấn mạnh đóng góp cho công ty.'),
        ('Kinh nghiệm làm việc', 'Sắp xếp theo thứ tự thời gian ngược (mới nhất ở trên). Mỗi mục ghi tên công ty, vị trí, thời gian. Mô tả trách nhiệm và thành tích, lượng hóa bằng số liệu cụ thể.'),
        ('Trình độ học vấn', 'Thứ tự thời gian ngược. Ghi tên trường, ngành, thời gian, GPA (nếu cao), học bổng, khóa học liên quan.'),
        ('Kỹ năng', 'Chia ba nhóm: chuyên môn (lập trình, phân tích dữ liệu...), mềm (giao tiếp, làm việc nhóm...), công cụ (Office, Git, Docker...).'),
        ('Dự án', 'Rất hữu ích cho sinh viên/ít kinh nghiệm. Ghi tiêu đề, thời gian, vai trò, mô tả, kỹ năng sử dụng, kết quả lượng hóa.'),
        ('Hoạt động ngoại khóa', 'Chọn hoạt động liên quan đến công việc ứng tuyển. Thể hiện sự năng động, kỹ năng mềm, tinh thần trách nhiệm.'),
        ('Sở thích', 'Chỉ ghi nếu liên quan đến công việc. Cụ thể hóa: "Đọc sách về AI" thay vì "Đọc sách".'),
        ('Người tham chiếu', 'Ghi tên, chức vụ, công ty, SĐT, email. Phải xin phép trước khi ghi tên.'),
        ('Đặt tên file', 'Format: "CV - Họ tên - Vị trí.pdf" hoặc "CV - Vị trí - Họ tên.pdf". Ưu tiên PDF để giữ layout.'),
    ]
    for i, (name, desc) in enumerate(cv_parts, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {name}: ')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run2 = p.add_run(desc)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(13)

    add_para(doc, '3. Tối ưu ATS', bold=True)
    add_para(doc,
        'ATS (Applicant Tracking System) là hệ thống lọc CV tự động. Để tối ưu ATS: '
        '(1) Sử dụng từ khóa từ JD trong mục tiêu, kinh nghiệm, kỹ năng; '
        '(2) Cấu trúc rõ ràng với tiêu đề chuẩn; '
        '(3) Định dạng Word hoặc PDF; '
        '(4) Không dùng bảng biểu phức tạp, hình ảnh, ký tự đặc biệt; '
        '(5) Font chữ chuẩn (Times New Roman, Arial).'
    )

    add_para(doc, '4. Chín lỗi thường gặp', bold=True)
    cv_errors = [
        'Email không chuyên nghiệp (nickname, ký tự đặc biệt)',
        'Mục tiêu nghề nghiệp chung chung, không cụ thể',
        'Mục tiêu quá dài dòng, không súc tích',
        'Mục tiêu không cá nhân hóa cho từng công ty',
        'Lỗi chính tả trong CV',
        'Không phân chia mục tiêu ngắn hạn/dài hạn',
        'Mục tiêu không thực tế, vượt quá khả năng',
        'Thời gian không theo trình tự ngược',
        'Thông tin thời gian không chính xác',
    ]
    for i, err in enumerate(cv_errors, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {err}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

    add_heading_styled(doc, 'II. Email xin việc', level=2)

    add_para(doc, '1. Định nghĩa và tình huống sử dụng', bold=True)
    add_para(doc,
        'Email xin việc là phương thức giao tiếp kỹ thuật số chính thức, được sử dụng để gửi CV '
        'và thư ứng tuyển đến nhà tuyển dụng. Email khác với cover letter ở chỗ: email là nội dung '
        'inline (trực tiếp trong thân email), ngắn gọn (150-200 từ), trong khi cover letter là file '
        'riêng chi tiết hơn (1 trang A4). Email được dùng khi nhà tuyển dụng không yêu cầu cover letter riêng.'
    )

    add_para(doc, '2. Bảy phần cấu trúc email', bold=True)
    email_parts = [
        ('Dòng chủ đề', 'Format: "Họ tên - Vị trí ứng tuyển - Công ty". Rõ ràng, ngắn gọn, chứa từ khóa.'),
        ('Người nhận', 'Email chính của phòng nhân sự hoặc người tuyển dụng.'),
        ('Lời chào', 'Tùy ngữ cảnh: "Kính gửi Quý Công ty" hoặc "Kính gửi Anh/Chị [Tên]".'),
        ('Nội dung (3 phần)', 'Giới thiệu (1-2 câu): họ tên, vị trí. Lý do + kỹ năng (2-3 câu). Kết luận: nhấn mạnh đóng góp cho công ty.'),
        ('Lời kết', '"Trân trọng," hoặc "Trân trọng cảm ơn,".'),
        ('File đính kèm', 'CV và cover letter (nếu có), định dạng PDF hoặc .docx. Đặt tên file chuẩn.'),
        ('Chữ ký', 'Họ tên, SĐT, email, LinkedIn/portfolio.'),
    ]
    for i, (name, desc) in enumerate(email_parts, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {name}: ')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run2 = p.add_run(desc)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(13)

    add_para(doc, '3. So sánh email xin việc và cover letter', bold=True)
    add_para(doc,
        'Email xin việc và cover letter có mục đích tương tự nhưng khác nhau về định dạng '
        'và mức độ chi tiết. Email là nội dung inline trong thân email, ngắn gọn (150-200 từ), '
        'dùng để giới thiệu và hướng dẫn đọc CV. Cover letter là file riêng (.docx/.pdf), '
        'chi tiết hơn (1 trang A4), giải thích sâu lý do ứng tuyển và sự phù hợp. '
        'Khi nhà tuyển dụng yêu cầu cover letter riêng, email chỉ cần ngắn gọn giới thiệu '
        'và liệt kê file đính kèm.'
    )

    add_para(doc, '4. Bảy lỗi thường gặp khi viết email', bold=True)
    email_errors = [
        'Hiểu nhầm email chỉ cần kỹ năng đơn giản — email cần chuyên nghiệp như cover letter',
        'Lạm dụng gửi quá nhiều email cho cùng một vị trí',
        'Định dạng kém: font, màu sắc, khoảng cách không phù hợp',
        'Nội dung dài dòng, không súc tích',
        'Lỗi chính tả — không kiểm tra trước khi gửi',
        'Thiếu dòng chủ đề (subject line)',
        'Quên file đính kèm khi đã nhắc đến CV trong nội dung',
    ]
    for i, err in enumerate(email_errors, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {err}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

def create_practice(doc):
    """Tạo phần thực hành"""
    doc.add_page_break()
    add_heading_styled(doc, 'PHẦN II: THỰC HÀNH SOẠN THẢO VĂN BẢN HÀNH CHÍNH, HỒ SƠ ỨNG TUYỂN', level=1)

    # ==================== NHIỆM VỤ 1 ====================
    add_heading_styled(doc, 'Nhiệm vụ 1: Bộ 03 văn bản hành chính', level=2)

    add_para(doc,
        'Dựa trên tình huống giả định về dự án "Nền tảng AI hỗ trợ bảo tồn và giảng dạy '
        'tiếng Jrai (mở rộng sang tiếng Bahnar) cho chuyển đổi số tỉnh Gia Lai" tại '
        'Học viện Công nghệ Bưu chính Viễn thông, nhóm thực hiện soạn thảo bộ 03 văn bản: '
        'tờ trình, công văn và thông báo.'
    )

    add_para(doc,
        'Ba văn bản này được trình bày trong file riêng biệt: '
        'VanBanHanhChinh_PTIT_Rewritten.docx. '
        'Tất cả đều đảm bảo đầy đủ 09 thành phần thể thức theo NĐ30/2020/NĐ-CP, '
        'đã được kiểm tra bằng công cụ check_nd30.py với kết quả 11/11 PASSED.'
    )

    add_para(doc, 'Tóm tắt nội dung ba văn bản:', bold=True)

    add_para(doc, '1.1. Tờ trình (Số: 01/TTr-PTIT)', bold=True)
    add_para(doc,
        'Nội dung: Xin phê duyệt kinh phí và kế hoạch thực hiện đề tài nghiên cứu '
        '"Nền tảng AI hỗ trợ bảo tồn và giảng dạy tiếng Jrai". '
        'Người ký: Trưởng khoa Công nghệ thông tin.'
    )

    add_para(doc, '1.2. Công văn (Số: 02/CV-PTIT)', bold=True)
    add_para(doc,
        'Nội dung: Gửi Sở Khoa học và Công nghệ tỉnh Gia Lai đề nghị hợp tác, '
        'hỗ trợ thu thập dữ liệu ngôn ngữ Jrai, hướng dẫn chính sách và đồng tài trợ. '
        'Người ký: Phó Giám đốc Học viện.'
    )

    add_para(doc, '1.3. Thông báo (Số: 03/TB-PTIT)', bold=True)
    add_para(doc,
        'Nội dung: Thông báo kế hoạch triển khai dự án với lộ trình 12 tháng '
        '(tháng 6/2025 - tháng 5/2026) đến các đơn vị và cá nhân liên quan. '
        'Người ký: Phó Giám đốc Học viện.'
    )

    # ==================== NHIỆM VỤ 2 ====================
    doc.add_page_break()
    add_heading_styled(doc, 'Nhiệm vụ 2: CV cá nhân', level=2)

    add_para(doc,
        'CV được soạn cho thành viên Phạm Tuấn Anh, ứng tuyển vị trí AI Engineer Intern. '
        'CV được trình bày trong file CV_PhamTuanAnh_AIEngineerIntern_Rewritten.docx, '
        'đảm bảo các yêu cầu: font Times New Roman, lề A4 theo quy định, '
        'đầy đủ các thành phần chính, tối ưu ATS, độ dài 1-2 trang A4.'
    )

    add_para(doc, 'Các thành phần chính của CV:', bold=True)

    cv_parts = [
        'Thông tin cá nhân: Họ tên, email, SĐT, địa chỉ, GitHub',
        'Mục tiêu nghề nghiệp: Ngắn hạn (AI Engineer) + dài hạn (nghiên cứu AI)',
        'Học vấn: PTIT, Kỹ thuật Phần mềm, GPA 3.47/4.00, IELTS 7.0',
        'Kỹ năng: Python, C++, SQL; PyTorch, Transformers, Hugging Face, OpenCV; FastAPI, Docker, Git',
        'Kinh nghiệm: Cộng tác viên nghiên cứu CV & Robotics tại PTIT',
        'Dự án: Tinh chỉnh ViT cho ảnh y khoa, nhận dạng giọng nói tiếng Việt, LSTM ước tính cảm biến',
        'Hoạt động ngoại khóa: Câu lạc bộ PAYT tại PTIT',
    ]
    for i, part in enumerate(cv_parts, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {part}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

    # ==================== NHIỆM VỤ 3 ====================
    doc.add_page_break()
    add_heading_styled(doc, 'Nhiệm vụ 3: Thư ứng tuyển và Email', level=2)

    add_para(doc,
        'Thư ứng tuyển và email được soạn dựa trên CV của Phạm Tuấn Anh, '
        'ứng tuyển vị trí AI Engineer Intern tại VinAI. '
        'Nội dung được trình bày trong file ThuUngTuyen_Email_PhamTuanAnh_VinAI.docx.'
    )

    add_para(doc, 'Tóm tắt nội dung:', bold=True)

    add_para(doc, 'Thư ứng tuyển (250-300 chữ):', bold=True)
    add_para(doc,
        'Giới thiệu bản thân là sinh viên PTIT, nêu lý do ứng tuyển vào VinAI. '
        'Trình bày kinh nghiệm nghiên cứu CV và NLP, các dự án đã hoàn thành (ViT, speech recognition, LSTM). '
        'Nhấn mạnh đóng góp dự kiến: phát triển mô hình AI, tối ưu pipeline ML, đóng góp vào sản phẩm AI. '
        'Kết thúc bằng lời cảm ơn và mong muốn phỏng vấn.'
    )

    add_para(doc, 'Email gửi hồ sơ:', bold=True)
    add_para(doc,
        'Chủ đề: "Phạm Tuấn Anh - AI Engineer Intern - VinAI". '
        'Nội dung 3 phần: giới thiệu ngắn gọn, lý do ứng tuyển + kỹ năng phù hợp, '
        'kết luận nhấn mạnh đóng góp cho công ty. '
        'File đính kèm: CV_PhamTuanAnh_AIEngineerIntern.pdf, ThuUngTuyen_PhamTuanAnh_AIEngineerIntern.pdf.'
    )

def create_conclusion(doc):
    """Tạo kết luận"""
    doc.add_page_break()
    add_heading_styled(doc, 'KẾT LUẬN', level=1)

    add_para(doc,
        'Qua quá trình thực hiện báo cáo, nhóm đã hoàn thành đầy đủ các nhiệm vụ '
        'theo yêu cầu của đề thi. Về lý thuyết, nhóm đã trình bày chi tiết chín thành phần '
        'thể thức chính của văn bản hành chính theo NĐ30/2020/NĐ-CP và các yêu cầu '
        'về CV cùng lưu ý khi soạn thảo email ứng tuyển.'
    )

    add_para(doc,
        'Về thực hành, nhóm đã soạn thảo thành công bộ 03 văn bản hành chính (tờ trình, '
        'công văn, thông báo) đảm bảo đúng thể thức, CV cá nhân chuyên nghiệp tối ưu ATS, '
        'thư ứng tuyển và email ứng tuyển đáp ứng yêu cầu. Tình huống giả định về dự án '
        'AI bảo tồn tiếng Jrai giúp nhóm vận dụng kiến thức vào thực tế.'
    )

    add_para(doc,
        'Báo cáo giúp nhóm hiểu sâu hơn về quy định thể thức văn bản hành chính, '
        'kỹ năng tạo lập CV và email chuyên nghiệp — những kỹ năng thiết yếu '
        'trong môi trường làm việc hiện đại.'
    )

def create_references(doc):
    """Tạo tài liệu tham khảo"""
    doc.add_page_break()
    add_heading_styled(doc, 'TÀI LIỆU THAM KHẢO', level=1)

    refs = [
        'Nghị định số 30/2020/NĐ-CP ngày 05 tháng 3 năm 2020 của Chính phủ về ban hành quy định về công tác văn thư.',
        'Phụ lục I, II, III Nghị định 30/2020/NĐ-CP về thể thức và kỹ thuật trình bày văn bản hành chính.',
        'Giáo trình Kỹ năng tạo lập văn bản tiếng Việt (SKD1103), Học viện Công nghệ Bưu chính Viễn thông.',
        'Hoàng Văn Hoa và Trần Thị Vân Hoa chủ biên (2012), Giáo trình giao tiếp trong kinh doanh, NxB Đại học Kinh tế Quốc dân.',
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'[{i}] {ref}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

def create_contribution_table(doc):
    """Tạo bảng đánh giá đóng góp"""
    doc.add_page_break()
    add_heading_styled(doc, 'BẢNG ĐÁNH GIÁ MỨC ĐỘ ĐÓNG GÓP CỦA TỪNG THÀNH VIÊN', level=1)

    # Tạo bảng
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ['STT', 'Họ và tên', 'Nhiệm vụ thực hiện', 'Mức độ đóng góp (%)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    # Dữ liệu mẫu
    data = [
        ['1', 'Phạm Tuấn Anh', 'Soạn thảo CV, thư ứng tuyển, email, tổng hợp báo cáo', '40%'],
        ['2', 'Thành viên 2', 'Soạn thảo tờ trình, công văn', '30%'],
        ['3', 'Thành viên 3', 'Soạn thảo thông báo, slide thuyết trình', '30%'],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

def create_slides_placeholder(doc):
    """Tạo placeholder cho slide thuyết trình"""
    doc.add_page_break()
    add_heading_styled(doc, 'SLIDE THUYẾT TRÌNH', level=1)

    add_para(doc,
        'Phần slide thuyết trình được trình bày dưới đây (2 slide/trang). '
        'Nội dung slide tóm tắt các nhiệm vụ đã thực hiện trong 7 phút trình bày.'
    )

    # Tạo bảng giả lập slide
    for slide_num in range(1, 7):
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        cell.height = Mm(80)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if slide_num == 1:
            run = p.add_run('SLIDE 1: GIỚI THIỆU\n\n')
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run2 = p.add_run('Đề tài: Nền tảng AI hỗ trợ bảo tồn và giảng dạy tiếng Jrai\n')
            run2.font.size = Pt(12)
            run2.font.name = 'Times New Roman'
            run3 = p.add_run('Nhóm: ... | Lớp: ...')
            run3.font.size = Pt(12)
            run3.font.name = 'Times New Roman'
        elif slide_num == 2:
            run = p.add_run('SLIDE 2: LÝ THUYẾT - 9 THÀNH PHẦN THỂ THỨC NĐ30\n\n')
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run2 = p.add_run('1. Quốc hiệu, tiêu ngữ\n2. Tên cơ quan chủ quản\n3. Tên cơ quan ban hành\n...')
            run2.font.size = Pt(12)
            run2.font.name = 'Times New Roman'
        elif slide_num == 3:
            run = p.add_run('SLIDE 3: LÝ THUYẾT - CV VÀ EMAIL\n\n')
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run2 = p.add_run('CV: 11 thành phần, tối ưu ATS\nEmail: Cấu trúc 7 phần, lỗi thường gặp')
            run2.font.size = Pt(12)
            run2.font.name = 'Times New Roman'
        elif slide_num == 4:
            run = p.add_run('SLIDE 4: THỰC HÀNH - 3 VĂN BẢN HÀNH CHÍNH\n\n')
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run2 = p.add_run('Tờ trình (01/TTr-PTIT)\nCông văn (02/CV-PTIT)\nThông báo (03/TB-PTIT)')
            run2.font.size = Pt(12)
            run2.font.name = 'Times New Roman'
        elif slide_num == 5:
            run = p.add_run('SLIDE 5: THỰC HÀNH - CV VÀ THƯ ỨNG TUYỂN\n\n')
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run2 = p.add_run('CV: Phạm Tuấn Anh - AI Engineer Intern\nThư ứng tuyển + Email: VinAI')
            run2.font.size = Pt(12)
            run2.font.name = 'Times New Roman'
        elif slide_num == 6:
            run = p.add_run('SLIDE 6: KẾT LUẬN\n\n')
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run2 = p.add_run('Hoàn thành đầy đủ các nhiệm vụ\nHiểu sâu về NĐ30 và kỹ năng tạo lập văn bản')
            run2.font.size = Pt(12)
            run2.font.name = 'Times New Roman'

        # Thêm khoảng cách giữa các slide
        if slide_num % 2 == 0 and slide_num < 6:
            doc.add_paragraph()

def main():
    doc = Document()

    # Thiết lập lề trang và khổ giấy A4
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)

    # Font mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # Tạo các phần
    create_cover_page(doc)
    doc.add_page_break()

    create_toc(doc)
    doc.add_page_break()

    create_foreword(doc)
    doc.add_page_break()

    create_theory(doc)
    doc.add_page_break()

    create_practice(doc)

    create_conclusion(doc)

    create_references(doc)

    create_contribution_table(doc)

    create_slides_placeholder(doc)

    # Lưu file
    output_path = os.path.join(OUTPUT_DIR, 'BCK_SKD1103_Lop_Nhom.docx')
    doc.save(output_path)
    print(f'Đã tạo báo cáo: {output_path}')

if __name__ == '__main__':
    main()
