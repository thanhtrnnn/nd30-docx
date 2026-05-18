#!/usr/bin/env python3
"""Create cover letter + email for VinAI application."""
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, size=14, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def add_para(doc, text, size=14, bold=False, italic=False,
             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size, bold, italic)
    return p

def add_page_break(doc):
    doc.add_page_break()

doc = Document()
s = doc.sections[0]
s.page_width = Mm(210)
s.page_height = Mm(297)
s.top_margin = Mm(20)
s.bottom_margin = Mm(20)
s.left_margin = Mm(30)
s.right_margin = Mm(15)

# ============================================================
# COVER LETTER (page 1)
# ============================================================
add_para(doc, "PHẠM TUẤN ANH", size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_para(doc, "Email: phamtuananh.ptit@gmail.com | ĐT: 0912 345 678", size=11, space_after=2)
add_para(doc, "Địa chỉ: Hà Nội | LinkedIn: linkedin.com/in/phamtuananh", size=11, space_after=12)

add_para(doc, "Ngày 18 tháng 05 năm 2026", size=12, italic=True, space_after=12)

add_para(doc, "Kính gửi: Phòng Tuyển dụng Nhân sự VinAI", size=13, bold=True, space_after=6)
add_para(doc, "Tôi viết thư này để ứng tuyển vị trí AI Engineer Intern tại VinAI. Với nền tảng kiến thức về trí tuệ nhân tạo, học sâu và xử lý ngôn ngữ tự nhiên, tôi mong muốn được đóng góp vào các nghiên cứu và sản phẩm AI tiên tiến của VinAI, đặc biệt trong lĩnh vực AI tiếng Việt và multi-modal AI.", size=13, space_after=6)

add_para(doc, "Trong quá trình học tập tại Học viện Công nghệ Bưu chính Viễn thông, tôi đã tích lũy được các kỹ năng chuyên môn:", size=13, space_after=4)

skills = [
    "PyTorch, Transformers, Hugging Face — xây dựng và fine-tuning mô hình AI",
    "Vision Transformer (ViT) fine-tuning cho phân loại hình ảnh",
    "Mô hình nhận dạng giọng nói (speech recognition) sử dụng LSTM",
    "Xử lý ngôn ngữ tự nhiên (NLP): tokenization, text classification, sentiment analysis",
    "Python, Git, Docker, Linux — công cụ phát triển và triển khai mô hình",
]
for skill in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Mm(1)
    run = p.add_run(f"• {skill}")
    set_font(run, 13)

add_para(doc, "Đóng góp dự kiến:", size=13, bold=True, space_after=4)
contribs = [
    "Phát triển và tối ưu các mô hình AI tiếng Việt, hỗ trợ VinAI trong nghiên cứu ngôn ngữ tự nhiên",
    "Đóng góp vào pipeline ML: data preprocessing, model training, evaluation, deployment",
    "Nghiên cứu và triển khai các kỹ thuật mới trong multi-modal AI (text-image-speech)",
    "Hỗ trợ đội ngũ kỹ sư trong việc benchmark và đánh giá hiệu suất mô hình",
]
for c in contribs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Mm(1)
    run = p.add_run(f"• {c}")
    set_font(run, 13)

add_para(doc, "Tôi tin rằng sự kết hợp giữa kiến thức chuyên môn về AI/ML và tinh thần học hỏi sẽ giúp tôi đóng góp tích cực vào các dự án của VinAI. Tôi rất mong có cơ hội được trao đổi thêm trong buổi phỏng vấn.", size=13, space_after=12)

add_para(doc, "Trân trọng,", size=13, space_after=24)
add_para(doc, "Phạm Tuấn Anh", size=13, bold=True, space_after=0)

# ============================================================
# EMAIL (page 2)
# ============================================================
add_page_break(doc)

add_para(doc, "EMAIL GỬI HỒ SƠ ỨNG TUYỂN", size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(doc, "Chủ đề: Phạm Tuấn Anh — AI Engineer Intern — VinAI", size=13, bold=True, space_after=8)

add_para(doc, "Kính gửi Anh/Chị Phòng Tuyển dụng Nhân sự VinAI,", size=13, space_after=6)

add_para(doc, "Em tên là Phạm Tuấn Anh, sinh viên năm cuối chuyên ngành Công nghệ Thông tin tại Học viện Công nghệ Bưu chính Viễn thông (PTIT). Em viết email này để ứng tuyển vị trí AI Engineer Intern tại VinAI.", size=13, space_after=6)

add_para(doc, "Qua quá trình học tập và thực hành, em đã xây dựng được nền tảng vững chắc về AI/ML, bao gồm: phát triển mô hình deep learning với PyTorch, fine-tuning Vision Transformer, xây dựng hệ thống nhận dạng giọng nói, và xử lý ngôn ngữ tự nhiên. Em đặc biệt quan tâm đến các nghiên cứu về AI tiếng Việt và multi-modal AI — những lĩnh vực trọng tâm của VinAI.", size=13, space_after=6)

add_para(doc, "Em tin rằng với kiến thức chuyên môn và tinh thần ham học hỏi, em có thể đóng góp tích cực vào các dự án nghiên cứu và phát triển sản phẩm AI của VinAI, đồng thời tích lũy được kinh nghiệm quý báu từ đội ngũ kỹ sư hàng đầu.", size=13, space_after=6)

add_para(doc, "Em xin đính kèm CV và thư ứng tuyển chi tiết. Rất mong nhận được phản hồi từ Anh/Chị.", size=13, space_after=6)

add_para(doc, "Em xin cảm ơn!", size=13, space_after=12)

add_para(doc, "Trân trọng,", size=13, space_after=18)
add_para(doc, "Phạm Tuấn Anh", size=13, bold=True, space_after=2)
add_para(doc, "SĐT: 0912 345 678", size=11, space_after=2)
add_para(doc, "Email: phamtuananh.ptit@gmail.com", size=11, space_after=2)
add_para(doc, "LinkedIn: linkedin.com/in/phamtuananh", size=11, space_after=0)

add_para(doc, "File đính kèm:", size=12, bold=True, space_after=4)
add_para(doc, "1. CV_PhamTuanAnh_AIEngineerIntern.pdf", size=11, space_after=2)
add_para(doc, "2. ThuUngTuyen_PhamTuanAnh_VinAI.pdf", size=11, space_after=0)

output_path = 'output/ThuUngTuyen_Email_PhamTuanAnh_VinAI.docx'
doc.save(output_path)
print(f'Da tao: {output_path}')
