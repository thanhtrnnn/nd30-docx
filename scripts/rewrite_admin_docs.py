#!/usr/bin/env python3
"""
Tạo lại 3 văn bản hành chính (Tờ trình, Công văn, Thông báo) từ đầu
với đúng format theo common.js và các generator scripts.
Spacing đã giảm: body=6pt (minimum spec), section headers/titles tighter.
"""
from docx import Document
from docx.shared import Pt, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# CONSTANTS từ common.js / generators
# ============================================================
PAGE_W = 11906  # A4 width DXA
PAGE_H = 16838  # A4 height DXA
MARGIN_TOP = 1134   # 20mm (NĐ30 minimum)
MARGIN_RIGHT = 851
MARGIN_BOTTOM = 1134 # 20mm
MARGIN_LEFT = 1701

COL_LEFT = 4082   # Header left col DXA (~45%)
COL_RIGHT = 4989  # Header right col DXA (~55%)

# Font sizes (half-points)
SZ_QH = 26       # Quốc hiệu: 13pt
SZ_TN = 28       # Tiêu ngữ: 14pt
SZ_CQ = 26       # Cơ quan: 13pt
SZ_SO = 26       # Số: 13pt
SZ_DIA_DANH = 28 # Địa danh: 14pt
SZ_TEN_LOAI = 28 # Tên loại: 14pt
SZ_BODY = 28      # Nội dung: 14pt
SZ_TRICH_YEU = 28 # Trích yếu: 14pt
SZ_TRICH_YEU_CV = 24 # Trích yếu CV V/v: 12pt
SZ_NOI_NHAN_TITLE = 24  # "Nơi nhận:" 12pt
SZ_NOI_NHAN_ITEM = 22   # Items: 11pt
SZ_CHUC_VU = 28   # Chức vụ: 14pt
SZ_CHU_KY = 28     # Chữ ký: 14pt

# Body spacing (from mau_to_trinh.docx template)
SP_BODY_AFTER = 120   # after=120 DXA (6pt)
SP_LINE = 276         # line=276 auto (body paragraphs)
SP_HEADER_LINE = 240  # line=240 auto (header/footer table paragraphs)

# Section/title spacing (tighter)
SP_SECTION_BEFORE = 30
SP_SECTION_AFTER = 0
SP_TITLE_BEFORE = 240   # template: 240
SP_TITLE_AFTER = 0      # template: 0
SP_SUBTITLE_BEFORE = 0  # template: 0
SP_SUBTITLE_AFTER = 0   # template: 0

INDENT_FIRST = 567  # 1cm (567 DXA)

# ============================================================
# HELPERS
# ============================================================

def set_run_font(run, size_half_pt, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_half_pt / 2)
    run.bold = bold
    run.italic = italic
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_para(container, text, size=SZ_BODY, bold=False, italic=False,
             alignment=WD_ALIGN_PARAGRAPH.LEFT, no_indent=False,
             spacing_before=None, spacing_after=None, line_spacing=None,
             first_line_indent=None):
    if hasattr(container, 'add_paragraph'):
        para = container.add_paragraph()
    else:
        if container.paragraphs:
            para = container.paragraphs[0] if not container.paragraphs[0].text else container.add_paragraph()
        else:
            para = container.add_paragraph()

    para.alignment = alignment

    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)

    spacing = OxmlElement('w:spacing')
    sb = spacing_before if spacing_before is not None else 0
    sa = spacing_after if spacing_after is not None else SP_BODY_AFTER
    sl = line_spacing if line_spacing is not None else SP_LINE
    spacing.set(qn('w:before'), str(sb))
    spacing.set(qn('w:after'), str(sa))
    spacing.set(qn('w:line'), str(sl))
    spacing.set(qn('w:lineRule'), 'auto')
    old_sp = pPr.find(qn('w:spacing'))
    if old_sp is not None:
        pPr.remove(old_sp)
    pPr.append(spacing)

    if first_line_indent is not None:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), str(first_line_indent))
    elif not no_indent:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), str(INDENT_FIRST))
    else:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '0')
    old_ind = pPr.find(qn('w:ind'))
    if old_ind is not None:
        pPr.remove(old_ind)
    pPr.append(ind)

    run = para.add_run(text)
    set_run_font(run, size, bold, italic)
    return para

def add_spacer(container, before=40, after=0):
    return add_para(container, '', no_indent=True,
                    spacing_before=before, spacing_after=after)

def add_separator(container, text_length=None, side='center'):
    """Add underscore separator line (1/3 to 1/2 of text width)."""
    if text_length:
        line_len = max(12, text_length // 3)
    else:
        line_len = 20
    line_text = "_" * line_len
    para = add_para(container, line_text, size=20,  # 10pt
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
                    spacing_before=20, spacing_after=60,
                    line_spacing=SP_HEADER_LINE)
    return para

def set_cell_no_borders(cell):
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcBorders = OxmlElement('w:tcBorders')
    for bn in ['top', 'bottom', 'left', 'right']:
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'FFFFFF')
        b.set(qn('w:space'), '0')
        tcBorders.append(b)
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(tcBorders)

def set_cell_width(cell, width_dxa):
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')

def set_table_width(table, width_pct=100):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(width_pct))
    tblW.set(qn('w:type'), 'pct')

# ============================================================
# HEADER TABLE
# ============================================================

def create_header_table(doc, co_quan_chu_quan, co_quan_ban_hanh, so_ky_hieu,
                        dia_danh="Hà Nội", ngay_thang=", ngày 05 tháng 05 năm 2025",
                        show_trich_yeu_cv=False, trich_yeu_cv=""):
    # Single-row table: số ký hiệu + ngày tháng in same cell as upper content
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, 100)

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]
    set_cell_width(left, COL_LEFT)
    set_cell_width(right, COL_RIGHT)
    set_cell_no_borders(left)
    set_cell_no_borders(right)

    # ── LEFT: cơ quan + separator + số ký hiệu ──
    if co_quan_chu_quan:
        add_para(left, co_quan_chu_quan, size=SZ_CQ,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
                 spacing_before=0, spacing_after=0, line_spacing=SP_HEADER_LINE)
    add_para(left, co_quan_ban_hanh, size=SZ_CQ, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=0, spacing_after=0, line_spacing=SP_HEADER_LINE)
    add_separator(left, text_length=len(co_quan_ban_hanh))

    # Số ký hiệu — right after separator, tight spacing
    add_para(left, so_ky_hieu, size=SZ_SO,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=6, spacing_after=0, line_spacing=SP_HEADER_LINE)

    if show_trich_yeu_cv and trich_yeu_cv:
        add_para(left, trich_yeu_cv, size=SZ_TRICH_YEU_CV,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
                 spacing_before=6, spacing_after=0, line_spacing=SP_HEADER_LINE)

    # ── RIGHT: quốc hiệu + tiêu ngữ + separator + ngày tháng ──
    add_para(right, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", size=SZ_QH, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=0, spacing_after=0, line_spacing=SP_HEADER_LINE)
    add_para(right, "Độc lập - Tự do - Hạnh phúc", size=SZ_TN, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=0, spacing_after=0, line_spacing=SP_HEADER_LINE)
    add_separator(right, text_length=30)

    # Ngày tháng — right after separator, tight spacing
    add_para(right, f"{dia_danh}{ngay_thang}", size=SZ_DIA_DANH, italic=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=6, spacing_after=0, line_spacing=SP_HEADER_LINE)

    return table

# ============================================================
# FOOTER TABLE
# ============================================================

def create_footer_table(doc, noi_nhan_items, chuc_vu_ky, nguoi_ky,
                        cap_ky="", chuc_vu_cap_tren=""):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, 100)

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]
    set_cell_width(left, 4535)   # 50%
    set_cell_width(right, 4535)  # 50%
    set_cell_no_borders(left)
    set_cell_no_borders(right)

    # Footer paragraphs: all before=0, after=0, line=240
    add_para(left, "Nơi nhận:", size=SZ_NOI_NHAN_TITLE, bold=True, italic=True,
             no_indent=True, spacing_before=0, spacing_after=0, line_spacing=SP_HEADER_LINE)
    for item in noi_nhan_items:
        add_para(left, item, size=SZ_NOI_NHAN_ITEM,
                 no_indent=True, spacing_before=0, spacing_after=0, line_spacing=SP_HEADER_LINE)

    valid_prefixes = ["KT", "TL", "TUQ", "TM", "Q"]
    if cap_ky.upper() in valid_prefixes and chuc_vu_cap_tren:
        add_para(right, f"{cap_ky.upper()}. {chuc_vu_cap_tren}", size=SZ_CHUC_VU,
                 bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
                 spacing_before=0, spacing_after=0, line_spacing=SP_HEADER_LINE)

    add_para(right, chuc_vu_ky, size=SZ_CHUC_VU, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=0, spacing_after=0, line_spacing=SP_HEADER_LINE)
    add_para(right, "(Ký, ghi rõ họ tên)", size=SZ_CHU_KY, italic=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=0, spacing_after=0, line_spacing=SP_HEADER_LINE)
    # 3 empty paragraphs for signature space (template: line=240 each = 720 DXA gap)
    for _ in range(3):
        add_para(right, "", size=SZ_CHU_KY,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
                 spacing_before=0, spacing_after=0, line_spacing=SP_HEADER_LINE)
    add_para(right, nguoi_ky, size=SZ_CHU_KY, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=0, spacing_after=0, line_spacing=SP_HEADER_LINE)

    return table

# ============================================================
# TỜ TRÌNH
# ============================================================

def create_to_trinh(doc):
    create_header_table(
        doc,
        co_quan_chu_quan="HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG",
        co_quan_ban_hanh="KHOA CÔNG NGHỆ THÔNG TIN 1",
        so_ky_hieu="Số: 01/2026/TTr-CNTT1",
        dia_danh="Hà Nội",
        ngay_thang=", ngày 10 tháng 05 năm 2026"
    )

    add_para(doc, "TỜ TRÌNH", size=SZ_TEN_LOAI, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_TITLE_BEFORE, spacing_after=SP_TITLE_AFTER)

    add_para(doc, "V/v đề nghị phê duyệt và cấp kinh phí đề tài khoa học và công nghệ",
             size=SZ_TRICH_YEU, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_SUBTITLE_BEFORE, spacing_after=SP_SUBTITLE_AFTER)
    add_separator(doc, text_length=55)

    add_para(doc, "Kính gửi: Ban Giám đốc Học viện Công nghệ Bưu chính Viễn thông",
             size=SZ_BODY, first_line_indent=567,
             spacing_before=240, spacing_after=60)

    add_para(doc, "Căn cứ Luật Khoa học và Công nghệ ngày 18/6/2013;",
             size=SZ_BODY, italic=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, "Căn cứ Nghị định số 30/2020/NĐ-CP ngày 05/3/2020 của Chính phủ về công tác văn thư;",
             size=SZ_BODY, italic=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, "Căn cứ Quy chế quản lý khoa học và công nghệ của Học viện Công nghệ Bưu chính Viễn thông;",
             size=SZ_BODY, italic=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "Khoa Công nghệ Thông tin 1 trình Ban Giám đốc Học viện nội dung đề xuất đề tài khoa học và công nghệ như sau:",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "I. SỰ CẦN THIẾT VÀ CĂN CỨ ĐỀ XUẤT",
             size=SZ_BODY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_SECTION_BEFORE, spacing_after=SP_SECTION_AFTER)
    add_para(doc, "Gia Lai là tỉnh có tỷ lệ đồng bào dân tộc thiểu số chiếm 44,7% dân số, trong đó người Jrai và Bahnar chiếm tỷ lệ lớn. Tiếng Jrai là ngôn ngữ được sử dụng rộng rãi trong sinh hoạt cộng đồng nhưng đang đối mặt với nguy cơ mai một do quá trình đô thị hóa và thiếu hụt tài liệu giảng dạy chuẩn hóa. Việc ứng dụng trí tuệ nhân tạo (AI) vào bảo tồn và giảng dạy tiếng Jrai không chỉ giúp lưu giữ di sản ngôn ngữ mà còn thúc đẩy chuyển đổi số trong giáo dục tại tỉnh Gia Lai.",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, "Căn cứ Nghị quyết số 57-NQ/TW ngày 22/12/2024 của Bộ Chính trị về đột phá phát triển khoa học, công nghệ, đổi mới sáng tạo và chuyển đổi số quốc gia; Căn cứ Quyết định số 1131/QĐ-BTTTT ngày 15/4/2025 của Bộ Thông tin và Truyền thông về Kế hoạch chuyển đổi số ngành TT&TT năm 2025, Khoa Công nghệ Thông tin 1 đề xuất thực hiện đề tài nghiên cứu ứng dụng AI phục vụ bảo tồn và giảng dạy tiếng Jrai.",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "II. NỘI DUNG ĐỀ TÀI",
             size=SZ_BODY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_SECTION_BEFORE, spacing_after=SP_SECTION_AFTER)
    for item in [
        "- Tên đề tài: Nền tảng AI hỗ trợ bảo tồn và dạy học tiếng Jrai (định hướng mở rộng tiếng Bahnar) phục vụ chuyển đổi số tỉnh Gia Lai.",
        "- Loại hình: Đề tài nghiên cứu ứng dụng và phát triển công nghệ.",
        "- Thời gian thực hiện: 12 tháng kể từ ngày được phê duyệt.",
        "- Đơn vị chủ trì: Khoa Công nghệ Thông tin 1 – Học viện Công nghệ Bưu chính Viễn thông.",
        "- Chủ nhiệm đề tài: TS. Phạm Vũ Minh Tú.",
    ]:
        add_para(doc, item, size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "III. DỰ KIẾN KINH PHÍ",
             size=SZ_BODY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_SECTION_BEFORE, spacing_after=SP_SECTION_AFTER)
    add_para(doc, "Tổng kinh phí đề nghị phê duyệt: 2.100.000.000 đồng (Hai tỷ một trăm triệu đồng) trong thời gian 12 tháng, từ nguồn ngân sách khoa học và công nghệ của Học viện.",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "IV. ĐỀ NGHỊ",
             size=SZ_BODY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_SECTION_BEFORE, spacing_after=SP_SECTION_AFTER)
    add_para(doc, "Khoa Công nghệ Thông tin 1 kính đề nghị Ban Giám đốc Học viện Công nghệ Bưu chính Viễn thông xem xét, phê duyệt đề tài khoa học và công nghệ nêu trên và cấp kinh phí để triển khai theo quy định.",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    create_footer_table(
        doc,
        noi_nhan_items=[
            "- Như kính gửi (để phê duyệt);",
            "- Phòng Khoa học, Công nghệ và HTQT (để phối hợp);",
            "- Lưu: VT, Khoa CNTT1."
        ],
        chuc_vu_ky="TRƯỞNG KHOA CÔNG NGHỆ THÔNG TIN 1",
        nguoi_ky="ThS. Nguyễn Duy Phương"
    )

# ============================================================
# CÔNG VĂN
# ============================================================

def create_cong_van(doc):
    create_header_table(
        doc,
        co_quan_chu_quan="",
        co_quan_ban_hanh="HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG",
        so_ky_hieu="Số: 01/2026/CV-HV",
        show_trich_yeu_cv=True,
        trich_yeu_cv="V/v đề nghị phối hợp triển khai đề tài khoa học và công nghệ",
        dia_danh="Hà Nội",
        ngay_thang=", ngày 15 tháng 05 năm 2026"
    )

    add_para(doc, "Kính gửi: Sở Khoa học và Công nghệ tỉnh Gia Lai",
             size=SZ_BODY, first_line_indent=567,
             spacing_before=240, spacing_after=60)

    paragraphs_text = [
        "Học viện Công nghệ Bưu chính Viễn thông (PTIT) đang triển khai đề tài nghiên cứu khoa học và công nghệ: \"Nền tảng AI hỗ trợ bảo tồn và dạy học tiếng Jrai (định hướng mở rộng tiếng Bahnar) phục vụ chuyển đổi số tỉnh Gia Lai\", do TS. Phạm Vũ Minh Tú làm chủ nhiệm, thuộc Khoa Công nghệ Thông tin 1.",
        "Đề tài bao gồm ba thành phần chính: (1) Nền tảng AI tìm kiếm và tạo nội dung số tiếng Jrai (từ điển ba ngôn ngữ Jrai-Việt-Anh, chatbot AI hỏi-đáp, công cụ chuyển văn bản thành giọng nói và ngược lại); (2) Hệ sinh thái E-learning với khóa học viết, đọc, nói tiếng Jrai ở 3 cấp độ; (3) Cổng kết nối cộng đồng mở rộng từ điển và bách khoa toàn thư văn hóa Jrai.",
        "Học viện Công nghệ Bưu chính Viễn thông trân trọng đề nghị Sở Khoa học và Công nghệ tỉnh Gia Lai xem xét phối hợp trong các nội dung sau:",
    ]
    for text in paragraphs_text:
        add_para(doc, text, size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    list_items = [
        "1. Hỗ trợ kết nối với các cơ sở giáo dục, cộng đồng dân tộc và cơ quan quản lý tại địa phương phục vụ thu thập ngữ liệu và thử nghiệm nền tảng.",
        "2. Tham gia góp ý định hướng chính sách và đánh giá kết quả thử nghiệm nền tảng AI tiếng Jrai tại tỉnh Gia Lai.",
        "3. Xem xét hỗ trợ kinh phí phối hợp trong khuôn khổ chương trình khoa học và công nghệ của tỉnh theo thẩm quyền.",
    ]
    for item in list_items:
        add_para(doc, item, size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "Học viện Công nghệ Bưu chính Viễn thông rất mong nhận được phản hồi từ Quý Sở. Đầu mối liên hệ: TS. Phạm Vũ Minh Tú – ĐT: 090 459 2738, Email: tuptm@ptit.edu.vn.",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    create_footer_table(
        doc,
        noi_nhan_items=[
            "- Như kính gửi;",
            "- Phòng KH,CN&HTQT (để theo dõi);",
            "- Lưu: VT, PTIT."
        ],
        cap_ky="KT",
        chuc_vu_cap_tren="GIÁM ĐỐC",
        chuc_vu_ky="PHÓ GIÁM ĐỐC",
        nguoi_ky="PGS.TS. Trần Quang Minh"
    )

# ============================================================
# THÔNG BÁO
# ============================================================

def _set_cell_spacing(para, before=40, after=40, line=240):
    """Set spacing on a paragraph inside a table cell."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    old_sp = pPr.find(qn('w:spacing'))
    if old_sp is not None:
        pPr.remove(old_sp)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def _cap_first(s):
    """Capitalize first letter of string."""
    if not s:
        return s
    return s[0].upper() + s[1:]

def add_table_grid(doc, headers, rows, font_size=12):
    """Add a Table Grid table with headers and data rows.
    Cells containing ';' are split into bullet points.
    Header row repeats on each new page (tblHeader)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    set_table_width(table, 100)

    # Header row — repeat on page break
    hdr_row = table.rows[0]
    trPr = hdr_row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)

    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h.upper())
        set_run_font(run, font_size * 2, bold=True)
        _set_cell_spacing(p, before=40, after=40, line=240)

    # Data rows — capitalize first letter of each item
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            items = [s.strip() for s in cell_text.split(';') if s.strip()]
            if len(items) <= 1:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(_cap_first(cell_text))
                set_run_font(run, font_size * 2, bold=False)
                _set_cell_spacing(p, before=40, after=40, line=240)
            else:
                for j, item in enumerate(items):
                    if j == 0:
                        p = cell.paragraphs[0]
                    else:
                        p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(f"- {_cap_first(item)}")
                    set_run_font(run, font_size * 2, bold=False)
                    _set_cell_spacing(p, before=20, after=20, line=240)

    return table


def create_thong_bao(doc):
    create_header_table(
        doc,
        co_quan_chu_quan="",
        co_quan_ban_hanh="HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG",
        so_ky_hieu="Số: 01/2026/TB-HV",
        dia_danh="Hà Nội",
        ngay_thang=", ngày 20 tháng 05 năm 2026"
    )

    add_para(doc, "THÔNG BÁO", size=SZ_TEN_LOAI, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_TITLE_BEFORE, spacing_after=SP_TITLE_AFTER)

    add_para(doc, "V/v Triển khai đề tài khoa học và công nghệ",
             size=SZ_TRICH_YEU, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_SUBTITLE_BEFORE, spacing_after=SP_SUBTITLE_AFTER)
    add_separator(doc, text_length=40)

    add_para(doc, "Căn cứ Tờ trình số 01/2026/TTr-CNTT1 ngày 10/5/2026 của Khoa Công nghệ Thông tin 1 về việc đề nghị phê duyệt và cấp kinh phí đề tài;",
             size=SZ_BODY, italic=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, "Căn cứ Kế hoạch khoa học và công nghệ năm 2026 của Học viện Công nghệ Bưu chính Viễn thông;",
             size=SZ_BODY, italic=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "Học viện Công nghệ Bưu chính Viễn thông thông báo về kế hoạch triển khai đề tài khoa học và công nghệ như sau:",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "1. Tên đề tài",
             size=SZ_BODY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_SECTION_BEFORE, spacing_after=SP_SECTION_AFTER)
    add_para(doc, "Nền tảng AI hỗ trợ bảo tồn và dạy học tiếng Jrai (định hướng mở rộng tiếng Bahnar) phục vụ chuyển đổi số tỉnh Gia Lai.",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "2. Đơn vị và cá nhân phụ trách",
             size=SZ_BODY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_SECTION_BEFORE, spacing_after=SP_SECTION_AFTER)
    add_para(doc, "- Đơn vị chủ trì: Khoa Công nghệ Thông tin 1 – Học viện Công nghệ Bưu chính Viễn thông.",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, "- Chủ nhiệm đề tài: TS. Phạm Vũ Minh Tú – ĐT: 090 459 2738.",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, "- Đơn vị phối hợp: Phòng Khoa học, Công nghệ và HTQT; Sở Khoa học và Công nghệ tỉnh Gia Lai.",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "3. Thời gian và kế hoạch triển khai",
             size=SZ_BODY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_SECTION_BEFORE, spacing_after=SP_SECTION_AFTER)
    add_para(doc, "Thời gian thực hiện: 12 tháng, từ tháng 6/2026 đến tháng 5/2027. Kế hoạch chi tiết như sau:",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_table_grid(doc,
        headers=["Giai đoạn", "Thời gian", "Hoạt động chính", "Sản phẩm"],
        rows=[
            ["1. Khảo sát", "Tháng 6–7/2026",
             "Khảo sát ngôn ngữ, văn hóa, nhu cầu cộng đồng Jrai; nghiên cứu dự án hiện có (YSJ, UET-TLU, VNeID Jrai)",
             "Báo cáo khảo sát; xác định pipeline xử lý"],
            ["2. Thu thập dữ liệu", "Tháng 8–10/2026",
             "Chuẩn hóa từ điển Jrai-Việt-Anh (~4.000 từ); xây dựng corpus song ngữ; thu thập âm thanh; phát triển Q&A chatbot",
             "Bộ dữ liệu từ điển; corpus song ngữ; 500 cặp Q&A"],
            ["3. Phát triển AI", "Tháng 11/2026–1/2027",
             "Tinh chỉnh PhoGPT; triển khai RAG cho chatbot; xây dựng API từ điển; áp dụng Active Learning",
             "Mô hình AI tiếng Jrai; API từ điển; portal đóng góp cộng đồng"],
            ["4. Thử nghiệm", "Tháng 2–4/2027",
             "Ra mắt beta; 3 hội thảo phản hồi (2 tại PTIT Gia Lai + 1 ĐH địa phương); triển khai E-learning + ứng dụng đa nền tảng",
             "Ứng dụng iOS/Android/Web; khóa học E-learning"],
            ["5. Nghiệm thu", "Tháng 5/2027",
             "Tổng kết; đánh giá chất lượng; báo cáo Sở KH&CN Gia Lai",
             "Hồ sơ nghiệm thu; báo cáo tổng kết"],
        ],
        font_size=12
    )

    # Spacer after table
    sp = doc.add_paragraph()
    _set_cell_spacing(sp, before=0, after=120, line=240)

    add_para(doc, "4. Kinh phí thực hiện",
             size=SZ_BODY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_SECTION_BEFORE, spacing_after=SP_SECTION_AFTER)
    add_para(doc, "Tổng kinh phí được phê duyệt: 2.100.000.000 đồng (Hai tỷ một trăm triệu đồng).",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "5. Yêu cầu",
             size=SZ_BODY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True,
             spacing_before=SP_SECTION_BEFORE, spacing_after=SP_SECTION_AFTER)
    add_para(doc, "Các đơn vị liên quan thuộc Học viện phối hợp chặt chẽ với Khoa Công nghệ Thông tin 1 trong suốt quá trình triển khai. Mọi thắc mắc xin liên hệ Phòng Khoa học, Công nghệ và HTQT.",
             size=SZ_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "./.", size=SZ_BODY, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT, no_indent=True,
             spacing_before=60, spacing_after=120)

    create_footer_table(
        doc,
        noi_nhan_items=[
            "- Các Khoa, Phòng, Ban liên quan (để thực hiện);",
            "- Sở KH&CN tỉnh Gia Lai (để phối hợp);",
            "- Lưu: VT, HV."
        ],
        cap_ky="KT",
        chuc_vu_cap_tren="GIÁM ĐỐC",
        chuc_vu_ky="PHÓ GIÁM ĐỐC",
        nguoi_ky="PGS.TS. Trần Quang Minh"
    )

# ============================================================
# MAIN
# ============================================================

def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Emu(PAGE_W * 635)
    section.page_height = Emu(PAGE_H * 635)
    section.top_margin = Emu(MARGIN_TOP * 635)
    section.bottom_margin = Emu(MARGIN_BOTTOM * 635)
    section.left_margin = Emu(MARGIN_LEFT * 635)
    section.right_margin = Emu(MARGIN_RIGHT * 635)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    create_to_trinh(doc)
    doc.add_page_break()

    create_cong_van(doc)
    doc.add_page_break()

    create_thong_bao(doc)

    output_path = 'output/VanBanHanhChinh_PTIT_Rewritten.docx'
    doc.save(output_path)
    print(f'Đã tạo: {output_path}')

if __name__ == '__main__':
    main()
