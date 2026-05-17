#!/usr/bin/env python3
"""Sửa thể thức 3 văn bản hành chính theo NĐ30"""
from docx import Document
from docx.shared import Pt, Mm, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def fix_date_format(text):
    """Sửa ngày tháng: tháng 5 → tháng 05"""
    import re
    # Sửa "tháng X" (X từ 1-9) thành "tháng 0X"
    text = re.sub(r'tháng (\d)(?!\d)', lambda m: f'tháng 0{m.group(1)}', text)
    return text

def set_cell_text(cell, text, bold=False, italic=False, size=Pt(13)):
    """Set text cho cell với font TNR"""
    # Xóa nội dung cũ
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''

    # Set nội dung mới
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.bold = bold
    run.italic = italic

    # Fix font East Asia
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    return para

def fix_to_trinh(doc):
    """Sửa tờ trình: header, ngày tháng"""
    # Table 0 là header table của tờ trình
    table = doc.tables[0]

    # Cột trái: Tên CQ chủ quản (KHÔNG đậm) + Tên CQ ban hành (ĐẬM) + gạch ngang
    left_cell = table.rows[0].cells[0]
    left_cell.text = ''
    p1 = left_cell.paragraphs[0]
    run1 = p1.add_run('BỘ THÔNG TIN VÀ TRUYỀN THÔNG')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(13)
    run1.bold = False  # KHÔNG đậm

    p2 = left_cell.add_paragraph()
    run2 = p2.add_run('HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(13)
    run2.bold = True  # ĐẬM

    p3 = left_cell.add_paragraph()
    run3 = p3.add_run('–––––––––––')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(13)
    run3.bold = False

    p4 = left_cell.add_paragraph()
    run4 = p4.add_run('Số: 01/2025/TTr-PTIT')
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(13)
    run4.bold = False

    # Cột phải: Quốc hiệu (ĐẬM) + Tiêu ngữ (ĐẬM) + gạch ngang + ngày tháng (nghiêng)
    right_cell = table.rows[0].cells[1]
    right_cell.text = ''
    p1 = right_cell.paragraphs[0]
    run1 = p1.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(13)
    run1.bold = True

    p2 = right_cell.add_paragraph()
    run2 = p2.add_run('Độc lập - Tự do - Hạnh phúc')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(13)
    run2.bold = True

    p3 = right_cell.add_paragraph()
    run3 = p3.add_run('')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(13)

    p4 = right_cell.add_paragraph()
    run4 = p4.add_run('Hà Nội, ngày 05 tháng 05 năm 2025')
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(13)
    run4.italic = True

    # Xóa paragraph "Số: 01/2025/TTr-PTIT" cũ (nếu có)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'Số: 01/2025/TTr-PTIT':
            # Xóa paragraph bằng cách xóa nội dung
            para.clear()
            break

    # Sửa "tháng 5" → "tháng 05" trong toàn bộ document
    for para in doc.paragraphs:
        for run in para.runs:
            if 'tháng 5' in run.text:
                run.text = run.text.replace('tháng 5', 'tháng 05')

    # Sửa trong tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if 'tháng 5' in run.text:
                            run.text = run.text.replace('tháng 5', 'tháng 05')

    return doc

def fix_thong_bao(doc):
    """Sửa thông báo: ngày tháng"""
    # Sửa "tháng 5" → "tháng 05" trong tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if 'tháng 5' in run.text:
                            run.text = run.text.replace('tháng 5', 'tháng 05')

    # Sửa trong paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if 'tháng 5' in run.text:
                run.text = run.text.replace('tháng 5', 'tháng 05')

    return doc

def main():
    input_path = 'output/VanBanHanhChinh_PTIT-Tờ-trình-Công-văn-Thông-báo.docx'
    output_path = 'output/VanBanHanhChinh_PTIT-To-trinh-Cong-van-Thong-bao_FIXED.docx'

    doc = Document(input_path)

    # Tìm vị trí bắt đầu của từng văn bản
    to_trinh_start = None
    cong_van_start = None
    thong_bao_start = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == 'TỜ TRÌNH':
            to_trinh_start = i
        elif text == 'CÔNG VĂN':
            cong_van_start = i
        elif text == 'THÔNG BÁO':
            thong_bao_start = i

    print(f'Tờ trình: paragraph {to_trinh_start}')
    print(f'Công văn: paragraph {cong_van_start}')
    print(f'Thông báo: paragraph {thong_bao_start}')

    # Sửa tờ trình (table 0 là header)
    doc = fix_to_trinh(doc)

    # Sửa thông báo (table 5 là header)
    doc = fix_thong_bao(doc)

    # Lưu file
    doc.save(output_path)
    print(f'Đã sửa: {output_path}')

if __name__ == '__main__':
    main()
